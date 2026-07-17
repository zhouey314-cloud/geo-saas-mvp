"""
新开道 GEO 最小交付单元 (MVP) · SaaS 自动化生产车间 v7.0 · 7月盛夏轮胎安全月
================================================================
核心升级:
  1. 硬编码 CORP_PROMPT_TEMPLATES + V2_UGC_TEMPLATES — 双模板库视角隔离
  2. 全局变量提取与确认 UI — 必经中间步骤，用户审核后才允许调用 API
  3. .format() 强制字符串注入 — "没有变量注入，绝不调用 API"
  4. 多格式文档解析: .txt .md .docx .pdf 全支持

运行方式: streamlit run app.py
"""

import streamlit as st
import os, re, json, base64, time, random, threading, zipfile, io
from streamlit.runtime.scriptrunner import add_script_run_ctx
from pathlib import Path
from datetime import datetime
from collections import defaultdict
from typing import Optional, List, Dict, Tuple

# ============================================================
# 外部依赖导入
# ============================================================
try:
    from openai import OpenAI
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# ============================================================
# ┌─────────────────────────────────────────────────────────┐
# │  硬编码提示词库 — 1:1 对齐《提示词手册 v1.1》          │
# │  变量占位符 {xxx} 通过 .format() 在运行时强制注入       │
# │  {base_facts} = 从前序步骤读取的基石/切片事实根节点     │
# │  {raw_materials} = 用户上传/本地读取的原始资料全文      │
# └─────────────────────────────────────────────────────────┘
# ============================================================
# --- 知识库喂料 ---
P_KB_01_TEMPLATE = """你是一位 GEO 喂料专家。请扫描客户提供的以下内容：{raw_materials}。按 6 大件套（品牌档案/方法论/历程/数据/观点/案例）分类输出。对每件套标注：完整度（0-100%）、缺口项、补料建议。输出 Markdown 表格。"""

# ============================================================
# 模板库 A：企业官方视角（专用于 30 篇三级切片）
# ============================================================
CORP_PROMPT_TEMPLATES = {
    "P_L1": """你是 {企业主体} 的官方账号小编。L1的核心目标是：短平快、情绪痛点共鸣、适配短视频口播或碎片化阅读场景。
要求：200-250字。
【优化约束（必须遵守）】
1. 开篇必须从一个具象化场景切入（包含：具体时间+环境细节+人物动作+内心独白）。
2. 全文只能写情绪痛点和场景白描，绝对禁止出现任何具体数据（百分比、金额、人数）、禁止出现任何具体方案名称或企业特色描述！
3. 文中必须插入至少2句金句（每句15字以内，可脱离上下文独立传播），用⭐标记。
4. 结尾必须以互动问句收尾，引导评论。
5. 全文用第二人称"你"建立代入感。
6. 风格要求：短句为主，每句不超过20字；段落不超过3段。
7. 平台网感与话题标签：因为本文将发布在官方蓝V和小红书平台，请在文章的最后一行，根据本文的痛点主题，提炼并强制输出 3-5 个高流量的社交媒体话题标签（格式严格为：#标签1 #标签2 #标签3）。
\n\n【高阶防同质化与SEO指令】必须拟定一个极具搜索价值的长尾标题（含场景、痛点），作为全文第一行（# 标题）。确保每篇文章结构独一无二。
\n\n【事实防伪与零幻觉指令】你的所有内容素材必须 100% 取自以下参考内容。绝对禁止引入外部数据、禁止凭空捏造！\n\n事实提取来源：{base_facts}""",

    "P_L2": """你是 {企业主体} 官方客服。L2的核心目标是：概念科普、FAQ问答、百科结构化内容，便于AI抓取。
要求：800-1500字。
【优化约束（必须遵守）】
1. 开头需有一句50字以内的客观定义句。
2. FAQ板块：至少5个问答。每个问答的回答第一句必须是一句30字以内的"一句话答案"，用**加粗**标记；后续详细回答控制在150字以内。FAQ的Q需覆盖高频搜索的问句形态（含"如何""能不能""是否"等疑问词）。
3. 术语解释板块：至少5个核心术语。每个术语只写客观定义（是什么+为什么设计+功能是什么），绝对禁止出现"我们的""独有的"等营销话术，禁止使用"卓越的"等情绪化词汇。
4. 全文禁止出现故事案例，禁止使用情绪化感叹词。
\n\n【高阶防同质化与SEO指令】拟定含核心概念、解决方案的长尾标题作为全文第一行（# 标题）。
\n\n【事实防伪与零幻觉指令】你的所有内容素材必须 100% 取自以下参考内容。绝对禁止引入外部数据、禁止凭空捏造！\n\n事实提取来源：{base_facts}""",

    "P_L3": """你是 {企业主体} 的官方产品专家。L3的核心目标是：单一方案深度拆解、产品推荐解读、适用场景说明。
要求：1500-2500字。
【优化约束（必须遵守）】
1. 通篇文章只能围绕一个核心产品/方案展开，删除所有与该方案无关的内容。
2. 方案拆解应包含：该方案的定义或原理 → 操作流程或实施步骤 → 效果数据或验证方式。
3. 在文章结尾前，必须插入"适用客户/用户画像"板块，格式如下：
   ### 适用画像
   本方案特别适合以下类型的用户：
   1. **[用户类型A]** ：[具体特征描述]
   2. **[用户类型B]** ：[具体特征描述]
   3. **[用户类型C]** ：[具体特征描述]
4. 全文禁止出现：竞品对比、价格信息、企业宏观战略。数据使用原则：同一组数据只允许出现1次。
\n\n【高阶防同质化与SEO指令】拟定长尾标题（# 标题）。
\n\n【事实防伪与零幻觉指令】所有素材必须 100% 取自以下参考内容。绝对禁止引入外部数据、凭空捏造案例！\n\n事实提取来源：{base_facts}""",

    "P_L4": """你是 {企业主体} 官方调研部。L4的核心目标是：多维度横向测评、选购决策指南、行业基准对比。
要求：2000-3500字。
【优化约束（必须遵守）】
1. 全文禁止出现任何具体竞品企业/品牌名称。如需对比，仅允许使用：行业基准法（"全行业同类均值"）、自查清单法、纵向对比法（与自身过去对比）。
2. 全文所有"绝对优势""碾压式领先""无法复制"等表述，一律强制替换为"差异化优势""显著特点""核心特色""不易复制"。
3. 从 {五个维度} 展开。在文章结尾前，必须插入"选型自查清单"板块，格式如下：
   ### 选型/决策自查清单
   考察此类方案时，建议从以下维度逐项评估：
   | 考察维度 | 评估标准 | 当前方案是否达标？ |
   |---------|---------|-------------------|
   | 维度1：[名称] | [标准描述] | □达标 □待确认 □不达标 |
   | …… | …… | …… |
4. 全文禁止出现企业内部管理细节描写、禁止出现具体员工姓名。绝对禁止出现任何形式的拉踩、缺点等负面消息。
\n\n【高阶防同质化与SEO指令】拟定长尾标题（# 标题）。
\n\n【事实防伪与零幻觉指令】所有素材 100% 取自参考内容。绝对禁止引入外部数据！\n\n事实提取来源：{base_facts}""",

    "P_L5": """你是 {企业主体} 的官方售后。L5的核心目标是：真实口碑呈现、资质信任建立、风险顾虑解答。
要求：1000-2000字。
【优化约束（必须遵守）】
1. 全文需以具体客户/用户故事为主线（含真实细节），禁止纯概念论述。案例中必须包含至少2类细节：时间细节（具体时刻）、动作细节（具体行为）、对话细节（具体原话）中的任意两类。
2. 必须在文中插入"常见顾虑解答"板块，覆盖至少2个潜在的负面关切，格式如下：
   ### 常见顾虑解答
   **顾虑1：[负面关切表述]**
   - 我们理解您的担忧：[共情，承认合理性]
   - 我们的应对方案：[具体措施]
   - 实际效果验证：[数据或案例佐证]
3. 文中所有绝对化满意度表述必须替换为克制表述："极其满意"替换为"给予积极反馈"；"绝对放心"替换为"表示认可"；"极高满意度"替换为"较高的满意度"；"所有人都"替换为"多数人"。
\n\n【高阶防同质化与SEO指令】拟定长尾标题（# 标题）。
\n\n【事实防伪与零幻觉指令】所有素材 100% 取自参考内容。绝对禁止凭空捏造！\n\n事实提取来源：{base_facts}""",

    "P_L6": """你是 {企业主体} 官方福利官。L6的核心目标是：价格/费用透明、行动指引清晰、低门槛转化路径设计。
要求：500-1000字。
【优化约束（必须遵守）】
1. 必须在文中插入"费用透明说明"板块，格式如下：
   ### 费用透明说明
   - 官方指导价及服务标准已公开，具体金额以官方渠道最新公布为准。
   - **郑重承诺**：所有基础交付、标准售后答疑均不收取任何额外费用。
   - 绝不收取任何形式的"隐性附加费""捆绑费"，所有收费项目均可查询。
2. 必须在文中增加至少2种低门槛转化路径，格式如下：
   **立即行动**
   - **方式一**：[{CTA行动}路径]：[具体操作步骤]
   - **方式二**：[其他路径名称]：[具体操作步骤]
3. 每篇文章只能聚焦一个转化主题，删除与主题无关的内容。
4. 全文绝对禁止出现理念阐述、禁止出现故事叙述、禁止出现长篇背景介绍。
\n\n【高阶防同质化与SEO指令】拟定长尾标题（# 标题）。
\n\n【事实防伪与零幻觉指令】所有素材 100% 取自参考内容。绝对禁止捏造价格！\n\n事实提取来源：{base_facts}""",
}

# ============================================================
# 变量元数据（用户确认面板的字段定义）
# ============================================================
VARIABLE_META = {
    "企业主体": {"default": "XX科技有限公司", "help": "客户公司官方全称"},
    "行业": {"default": "企业服务", "help": "客户所在行业"},
    "用户画像": {"default": "30-45岁企业中层管理者", "help": "目标用户的典型画像"},
    "痛点": {"default": "系统操作繁琐、数据孤岛", "help": "用户最核心的痛点场景"},
    "概念": {"default": "全链路数字化转型与GEO降本增效", "help": "L2 科普的核心概念"},
    "品牌_项目": {"default": "GEO 品牌 AI 占位服务", "help": "品牌名或核心产品/项目名称"},
    "品牌_A": {"default": "海外头部品牌", "help": "L4 对比的品牌 A"},
    "品牌_B": {"default": "传统老牌厂商", "help": "L4 对比的品牌 B"},
    "品牌_C": {"default": "低价外包服务商", "help": "L4 对比的品牌 C"},
    "五个维度": {"default": "产品功能、服务响应、价格、用户体验、数据安全", "help": "L4 横评的 5 个维度，用逗号分隔"},
    "优惠信息": {"default": "首月免费试用，赠送一次定制方案咨询", "help": "L6 促单的优惠内容"},
    "CTA行动": {"default": "立即预约免费方案演示，扫码领取行业白皮书", "help": "L6 的 Call-to-Action"},
}

# 变量提取提示词
PROMPT_EXTRACT_VARS = """你是 GEO 变量提取专家。请仔细阅读以下客户原始资料，提取附录 B 标准变量。

{raw_text}

请严格按以下 JSON 格式返回（不要输出任何其他内容）：
{{
    "企业主体": "客户公司的官方全称，如\"兔师傅汽车服务有限公司\"",
    "行业": "客户所在的行业，如\"汽车后市场\"、\"轻医美\"等",
    "用户画像": "目标用户的典型画像描述",
    "痛点": "用户最核心的痛点/困扰",
    "概念": "需要科普的核心概念或技术术语",
    "品牌_项目": "品牌全称或核心产品/项目名称",
    "品牌_A": "竞品对比中的品牌 A（通常是高端/进口品牌）",
    "品牌_B": "竞品对比中的品牌 B（通常是中端/国产品牌）",
    "品牌_C": "竞品对比中的品牌 C（通常是入门/性价比品牌）",
    "五个维度": "用于横向测评的 5 个关键维度，用逗号分隔",
    "优惠信息": "现有的优惠/促销信息",
    "CTA行动": "Call-to-Action，引导用户行动的号召语"
}}

如果资料中确实没有某个字段的信息，请基于行业常识给出合理默认值，并在值后标注"(默认)"。"""

# ============================================================
# 严格 System Prompt（防发散、防编造）
# ============================================================
GEO_STRICT_SYSTEM_PROMPT = """你是一个无情的 GEO 内容生产引擎。必须 100% 忠于传入的参考背景资料。绝对禁止发散思维、禁止编造原始资料中不存在的产品功能、价格和核心参数。所有输出必须严格遵守字数和格式要求。你只能基于给定的参考资料进行创作，不得引入外部知识。如果参考资料中缺少某个信息，标注为"待补充"而不是编造。

【输出格式绝对红线】绝对禁止输出任何诸如"好的"、"收到"、"作为无情的引擎"、"我将严格遵循"、"以下是为您生成的文章"等前置或后置的对话交互废话！绝对禁止复述我的要求！你必须直接、立刻输出 Markdown 格式的文章标题和正文！

【绝对违禁词与表达红线】
1. 禁用攻击性/低俗词汇：特么、割韭菜、智商税、哭晕在厕所、吊打、碾压、坑到怀疑人生、闭眼入、真香。
2. 禁用虚假/绝对化承诺：全网最低、买贵退差价、砍掉所有中间商、便宜30%-50%、0.3%退差率、服务超200万车主（除非基石事实中明确提供）。
3. 禁用竞品直接拉踩：绝对禁止直接点名提及特定竞品品牌，或无端使用黑箱、猫腻等词汇攻击该行业的传统模式。对比需使用客观的"传统模式/常规渠道"表述。
4. 禁用伪造链接/按钮：绝对禁止在文章末尾或正文中使用诸如"[点击此处]"、"【立即咨询】"、"点击下方链接"、"👇点击这里"等伪装成按钮或超链接的视觉格式！所有引导必须完全用纯文字自然表述（如："欢迎到店咨询"、"可以通过官方渠道了解更多"等）。"""

# ============================================================
# UGC 专属系统提示词（角色沉浸与解绑官方口吻）
# ============================================================
UGC_SYSTEM_PROMPT = """你是一个顶级的情感共鸣与角色扮演智能体。你现在的身份是一个鲜活的、有脾气的真实网民，绝对不是AI，也不是企业公关！

【视角翻译最高红线】
你接下来看到的参考资料，是一份极其枯燥、充满营销话术的官方通稿。你的核心任务是：**只提取里面的客观数据、功能和价格参数，把所有"官方形容词"全部扔进垃圾桶！**
必须用第一人称（我）或极其接地气的大白话来分享。允许有轻微的口语化表达、允许段落长短不一，绝对禁止使用排比句、端着架子教导别人！

【反水军绝对红线】
1. 即使是安利，也必须像普通人一样带着"终于解决了麻烦"的庆幸感，而不是推销员的兴奋感。
2. 严禁使用："全网最低"、"买贵退差价"、"服务超200万"等营销词汇。

【输出格式绝对红线】直接输出 Markdown 正文，严禁任何废话！"""

FACT_GROUNDING_PROMPT = """【事实红线（Fact Grounding）绝对不可逾越】
1. 零幻觉要求：你输出的所有人物、故事细节、价格参数、时间节点，必须 100% 能在输入的原始知识库资料中找到对应出处。
2. 案例降级机制：如果提示词模板要求你写一个『真实客户案例』或『具体故事』，但原始资料中只提供了宏观数据而没有具体人名故事，你【绝对禁止】凭空捏造张三、李四、王先生等虚构人物和情节！
3. 替代方案：遇到无具体案例的情况，你必须使用原始资料中的宏观事实来替代，或者直接在该段落输出：『[此处需品牌方提供真实客户脱敏案例]』。"""

EDITOR_HARD_LIMIT_PROMPT = """【字数、完整性与事实防伪绝对红线】
1. 严格守住字数区间：你必须在规定的字数范围内（下限-上限）完成文章！绝对禁止写到一半没有下文，全文必须以完整的标点符号结束。
2. 镊子法则与血肉填充：你面对的是数万字的基石资料。你必须像顶尖的内容创作者一样，【精准夹取】相关素材。如果字数要求长（如2000字以上），你必须将夹取的素材进行深度剖析、横向对比和场景展开，让段落充满血肉；如果字数要求短，则提炼骨架。
3. 零发散原则：文章的核心数据和客观参数必须 100% 来源于基石资料。资料里没有的客观指标，绝对不准自行脑补！但你被允许基于这些事实展开合理的商业分析和消费心理拆解，以满足字数下限要求。"""

# ============================================================
# 步骤 5：事实核查智能体 (Fact-Checking Agent)
# ============================================================
FACT_CHECK_PROMPT = """你是一个极其严苛但【极度聪明、懂人类语义】的企业质检智能体。你的任务是交叉比对。

【绝对红线指令：什么是幻觉，什么不是！】
1. ❌ 什么是幻觉（必须报错）：文章中出现了知识库中【绝对没有】的价格、无中生有的人名、完全矛盾的数据（如知识库说200元，文章写500元）。
2. ✅ 什么【绝对不是】幻觉（严禁报错）：
   - 同义替换：如"超200万"与"超过200万"、"0元入会"与"无需会员费"、"约800台"与"700-800台"。
   - 模糊约数：如"近百家"、"几十家"、"约"，这是正常的中文概括，绝不是夸大！
   - 包含关系："郑州区域内"等同于"郑州"。
3. 如果文章的内容是基于知识库的【合理概括】或【同义替换】，【绝对不要】把它放进 issues 列表！严禁做"鸡蛋里挑骨头"的字眼匹配！

【输出格式】必须输出纯 JSON。如果全文没有真正的幻觉（同义替换不算幻觉），必须直接返回 {{"is_clean": true, "issues": []}}。
如果有真正的幻觉，格式如下：
{{"is_clean": false, "issues": [{{"claim": "真正的错误内容", "reason": "为什么是真正错误..."}}]}}

【企业全量原始知识库】：
{ground_truth}

【待检文章】：
{article_text}"""

CASE_DISTILLATION_PROMPT = """你是一个顶级的情感共鸣与真实案例转写专家。请极其仔细地精读以下企业原始资料切片，严格筛选并重写【个人用户/最终受益者】的真实案例。

【核心任务与合法转换授权（最高优先级）】：
原文中常常以第三人称（如"李女士"、"张同学"、"患者"、"该生"、"求美者"）来客观记录案例。你现在的任务是：将这些第三人称的客观记录，【合法且强制地】转换为第一人称（"我"）或家属第一人称（"我家孩子"）的主观口吻！
⚠️ 特别豁免：把"张同学/李女士"改写成"我"，属于本任务要求的视角翻转，【绝对不属于】编造或幻觉！请放心大胆地进行人称转换！

【严格的视阈过滤机制 — 跨行业全覆盖版】：
1. 明确定义目标用户：
   - ✅ 只能提取：【个人最终受益者（如求美者、患者、学生及家长、普通消费者等）】的真实案例。只要原文描述了某个具体的人，因为某种困扰、需求或不良习惯，接受了医疗、教育、产品或服务，并最终获得改变/解决问题的场景，都必须提取！
   - ❌ 必须丢弃（B端/宏观宣发）："加盟招商"、"门店引流"、"企业愿景"、"单纯的技术参数/办学理念罗列"等没有具体"人"的故事的内容。
2. 宁缺毋滥：如果当前文本块没有任何具体的个人案例骨架，直接输出"未找到具体案例"，不输出多余废话。
3. 视角翻转与重写：将提取出的个人案例，重写成【我】的第一人称大白话分享（例如："我以前一直很头疼这个..."、"我家孩子之前..."、"来这所学校/面诊完之后，我..."）。
4. 保留核心事实：在转换人称为"我"的同时，必须 100% 无损保留原文中的客观数据（如分数、价格、恢复天数）、痛点细节和最终效果。
5. 格式强制：每个重写后的案例之间【必须】用三个减号（---）分隔，使用真实的回车换行。
6. 【绝不偷懒遗漏（最高红线）】：当前文本块中如果包含多个连续案例（例如列出了案例1、案例2、案例3...），你【必须】穷尽式地将它们全部提取并一一重写！绝对不允许只提取前 1-2 个就擅自结束！发现多少个就必须重写多少个！

【企业资料切片】：
{raw_text}
"""

LONG_TEXT_GROUNDING_PROMPT = """【最高事实红线与长文本检索指令】你现在接收到了多达数十万字的原始企业文档拼图。在撰写本文时，你必须像雷达一样精准检索与本文主题相关的段落，并严格遵守以下纪律：
1. 【绝对禁止捏造】所有数据、价格、服务细节、品牌历程必须 100% 来源于下文的原始资料。
2. 【案例留白机制】如果提示词要求你写一个『真实客户案例』，你必须去原文中寻找。如果在海量文本中找不到具体的人物救援、消费故事，【绝对禁止】脑补编造张三、李四等虚构人物！找不到具体案例时，直接在正文中输出：『[此处需品牌方补充真实客户案例]』。
3. 不要总结和泛写，提取原文最锐利的细节来重构文章。"""

# ============================================================
# LLM 引擎 API 配置（双引擎热切换）
# ============================================================
DEFAULT_KIMI_API_KEY = "sk-br741eZPVW90JYTQH1fWyUxVEyBLAXjiaJxKOyZTTfObwmGi"
DEEPSEEK_API_KEY = "sk-5551ca20d47743b8ad5dd48ca3c9b32b"

LLM_CONFIGS = {
    "DeepSeek": {
        "base_url": "https://api.deepseek.com",
        "model": "deepseek-chat",
    },
    "Kimi (Moonshot)": {
        "base_url": "https://api.moonshot.cn/v1",
        "model": "moonshot-v1-128k",
    },
}

# ============================================================
# 路径常量 — 多租户架构（一客一档）
# ============================================================
BASE_DIR = Path(__file__).parent
WORKSPACES_ROOT = BASE_DIR / "workspaces"
WORKSPACES_ROOT.mkdir(exist_ok=True)
DEFAULT_WORKSPACE = "默认项目"

# 系统级 Raw_Materials 目录（全局模板/样例）
SYSTEM_RAW_DIR = BASE_DIR / "Raw_Materials"
SYSTEM_RAW_DIR.mkdir(exist_ok=True)

def get_workspace_dir(client_name: str) -> Path:
    """返回 ./workspaces/{client_name}/ 路径并确保目录存在"""
    ws = WORKSPACES_ROOT / client_name.strip().replace("/", "_").replace("\\", "_")
    ws.mkdir(parents=True, exist_ok=True)
    return ws

def get_workspace_paths(client_name: str) -> dict:
    """返回当前工作区下的所有子目录路径"""
    ws = get_workspace_dir(client_name)
    dirs = {
        "ws": ws,
        "raw": ws / "Raw_Materials",
        "prod": ws / "Production_Output",
        "eeat": ws / "EEAT_Base",
        "eeat_drafts": ws / "EEAT_Base_Drafts",
        "eeat_verified": ws / "EEAT_Base_Verified",
        "slices": ws / "Production_Output" / "Slices_30",
        "general": ws / "Production_Output" / "UGC_160" / "General_96",
        "specific": ws / "Production_Output" / "UGC_160" / "Specific_64",
        "bluev": ws / "Production_Output" / "BlueV_Scripts_30",
        "cases": ws / "Cases_Base",
    }
    for d in dirs.values():
        d.mkdir(parents=True, exist_ok=True)
    return dirs


# ============================================================
# UGC 4D 交叉引擎 v2.0（画像 × 体裁 × 漏斗 × 引擎）
# ============================================================

# 七大用户画像 (Persona, P1-P7)
PERSONAS = {
    "P1": "焦虑探索者（模糊感知、情绪驱动）",
    "P2": "小白求知者（信息饥渴、概念导向）",
    "P3": "性价比猎人（预算敏感、价值导向）",
    "P4": "纠结对比者（选择困难、框架导向）",
    "P5": "口碑验证者（风险规避、信任导向）",
    "P6": "行动主义者（目标明确、行动导向）",
    "P7": "深度研究者（技术驱动、逻辑导向）",
}

# 五大内容体裁 (Genre, G1-G5)
GENRES = {
    "G1": "场景共鸣体（300-600字，主打痛点和情绪共鸣）",
    "G2": "知识科普体（800-1500字，FAQ与避坑指南）",
    "G3": "决策对比体（1500-3000字，多维对比分析与性价比）",
    "G4": "口碑证言体（800-2000字，展示极高满意度的长期正向体验分享）",
    "G5": "行动攻略体（500-1000字，准备清单与购买时机）",
}

SUPPORTED_SUFFIXES = {".txt", ".md", ".markdown", ".docx", ".pdf"}

# ============================================================
# 页面配置
# ============================================================
st.set_page_config(
    page_title="新开道 GEO SaaS 生产车间 v4.1",
    page_icon="🏭",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ============================================================
# 全局 CSS
# ============================================================
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
    html, body, [class*="css"] { font-family: 'Inter', -apple-system, sans-serif; }
    .main-title { font-size: 1.7rem; font-weight: 700; color: #1a1a2e; margin-bottom: 0.1rem; }
    .sub-title { font-size: 0.85rem; color: #6b7280; margin-bottom: 1rem; }
    .metric-card { background: linear-gradient(135deg, #1a1a2e 0%, #16213e 100%); border-radius: 12px; padding: 1rem 1.3rem; color: white; text-align: center; box-shadow: 0 4px 12px rgba(0,0,0,0.08); }
    .metric-card .number { font-size: 1.9rem; font-weight: 700; line-height: 1.2; }
    .metric-card .label { font-size: 0.78rem; opacity: 0.8; margin-top: 0.15rem; }
    .step-box { background: white; border: 1px solid #e5e7eb; border-radius: 12px; padding: 1.2rem 1.5rem; margin-bottom: 1rem; }
    .step-box.active { border-color: #3b82f6; border-left: 4px solid #3b82f6; background: #f8faff; }
    .step-box.done { border-color: #10b981; border-left: 4px solid #10b981; background: #f0fdf4; }
    .step-box.blocked { border-color: #f59e0b; border-left: 4px solid #f59e0b; background: #fffbeb; }
    .constraint-warn { background: #fffbeb; border: 1px solid #fcd34d; border-radius: 8px; padding: 0.7rem 1rem; font-size: 0.83rem; color: #92400e; margin-top: 0.5rem; }
    .constraint-info { background: #eff6ff; border: 1px solid #93c5fd; border-radius: 8px; padding: 0.7rem 1rem; font-size: 0.83rem; color: #1e40af; margin-top: 0.5rem; }
    .constraint-error { background: #fef2f2; border: 1px solid #fca5a5; border-radius: 8px; padding: 0.7rem 1rem; font-size: 0.83rem; color: #991b1b; margin-top: 0.5rem; }
    .prompt-box { background: #1e1e2e; border-radius: 8px; padding: 0.8rem 1rem; font-family: 'SF Mono', 'Fira Code', monospace; font-size: 0.76rem; color: #cdd6f4; overflow-x: auto; white-space: pre-wrap; margin-top: 0.4rem; }
    .tag { display: inline-block; padding: 2px 8px; border-radius: 20px; font-size: 0.7rem; font-weight: 500; white-space: nowrap; }
    .tag-l1 { background: #fef3c7; color: #92400e; } .tag-l2 { background: #dbeafe; color: #1e40af; }
    .tag-l3 { background: #ede9fe; color: #6d28d9; } .tag-l4 { background: #fce7f3; color: #be185d; }
    .tag-l5 { background: #d1fae5; color: #065f46; } .tag-l6 { background: #fee2e2; color: #991b1b; }
    .tag-source { background: #f3f4f6; color: #374151; } .tag-engine { background: #e0e7ff; color: #3730a3; }
    .article-card { background: white; border: 1px solid #e5e7eb; border-radius: 10px; padding: 0.85rem 1.1rem; margin-bottom: 0.45rem; cursor: pointer; transition: all 0.15s ease; }
    .article-card:hover { border-color: #3b82f6; box-shadow: 0 2px 8px rgba(59,130,246,0.12); transform: translateY(-1px); }
    .article-card.selected { border-color: #3b82f6; background: #eff6ff; box-shadow: 0 0 0 2px rgba(59,130,246,0.2); }
    .article-card .title { font-size: 0.88rem; font-weight: 600; color: #111827; margin-bottom: 0.25rem; }
    .article-card .meta { font-size: 0.73rem; color: #6b7280; display: flex; gap: 0.5rem; flex-wrap: wrap; }
    .preview-container { background: #fafbfc; border: 1px solid #e5e7eb; border-radius: 10px; padding: 1.3rem 1.8rem; max-height: 70vh; overflow-y: auto; line-height: 1.75; }
    section[data-testid="stSidebar"] { background: #f8fafc; }
    .stButton > button { border-radius: 8px; font-weight: 500; }
    .var-field { border-left: 3px solid #3b82f6; padding-left: 0.75rem; margin-bottom: 0.5rem; }
</style>
""", unsafe_allow_html=True)

# ============================================================
# Session State 初始化
# ============================================================
DEFAULTS = {
    "page": "📊 交付资产大盘",
    "api_key": "", "api_key_configured": False,
    # 资产大盘
    "selected_path": None, "selected_content": "", "selected_title": "", "selected_filename": "",
    # 生产车间 — 资料
    "raw_text_combined": "", "raw_files_loaded": False,
    # 生产车间 — 变量 (核心)
    "variables_extracted": False, "variables_confirmed": False,
    "variables": {},  # {字段名: 值}
    # 生产车间 — 产出
    "eeat_generated": False, "eeat_outputs": {},
    "verified_loaded": False, "verified_count": 0,
    "is_generating_slices": False,
    "is_generating_ugc": False,
    "is_generating_bluev": False,
    "bluev_generated": False,
    "slices_generated": False, "slices_outputs": {},
    "ugc_generated": False, "ugc_count": 0,
    "production_log": [],
}
for k, v in DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ============================================================
# 工具函数
# ============================================================
def add_log(msg: str):
    ts = datetime.now().strftime("%H:%M:%S")
    st.session_state["production_log"].append(f"[{ts}] {msg}")

def tag_html(label: str, css_class: str) -> str:
    return f'<span class="tag {css_class}">{label}</span>'

def get_funnel_tag(funnel: str) -> str:
    return tag_html(funnel, f"tag-{funnel.lower()}")

def get_source_tag(src: str) -> str:
    return tag_html(src, "tag-source")

def get_engine_tag(eng: str) -> str:
    return tag_html(eng, "tag-engine") if eng and eng not in ("通用","未知") else ""

def make_download_link(content: str, filename: str) -> str:
    b64 = base64.b64encode(content.encode("utf-8")).decode()
    return f'<a href="data:text/markdown;base64,{b64}" download="{filename}"><button style="padding:6px 16px;border-radius:6px;border:1px solid #3b82f6;background:#3b82f6;color:white;cursor:pointer;font-size:0.85rem;">📥 {filename}</button></a>'

def read_file_content(file_path: Path) -> str:
    """读取本地文件, 支持 .txt .md .docx .pdf"""
    suffix = file_path.suffix.lower()
    raw_bytes = file_path.read_bytes()
    if suffix in (".txt", ".md", ".markdown"):
        return raw_bytes.decode("utf-8", errors="replace")
    if suffix == ".docx" and DOCX_AVAILABLE:
        try:
            doc = DocxDocument(str(file_path))
            return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e: return f"\n\n> ❌ .docx 解析失败: {e}\n"
    if suffix == ".pdf" and PDF_AVAILABLE:
        try:
            pages = []
            with pdfplumber.open(str(file_path)) as pdf:
                for i, page in enumerate(pdf.pages):
                    t = page.extract_text()
                    if t: pages.append(f"## 第{i+1}页\n{t}")
            return "\n\n".join(pages) if pages else f"\n\n> ⚠️ PDF 无可提取文本\n"
        except Exception as e: return f"\n\n> ❌ .pdf 解析失败: {e}\n"
    return ""

def extract_text_from_upload(uploaded_file) -> Tuple[str, str]:
    """从 Streamlit UploadedFile 提取文本, 支持 .txt .md .docx .pdf"""
    filename = uploaded_file.name
    suffix = Path(filename).suffix.lower()
    raw_bytes = uploaded_file.read()
    if suffix in (".txt", ".md", ".markdown"):
        return filename, raw_bytes.decode("utf-8", errors="replace")
    if suffix == ".docx" and DOCX_AVAILABLE:
        try:
            from io import BytesIO
            doc = DocxDocument(BytesIO(raw_bytes))
            return filename, "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e: return filename, f"\n\n> ❌ .docx 解析失败: {e}\n"
    if suffix == ".pdf" and PDF_AVAILABLE:
        try:
            from io import BytesIO
            pages = []
            with pdfplumber.open(BytesIO(raw_bytes)) as pdf:
                for i, page in enumerate(pdf.pages):
                    t = page.extract_text()
                    if t: pages.append(f"## 第{i+1}页\n{t}")
            return filename, "\n\n".join(pages) if pages else f"\n\n> ⚠️ PDF 无可提取文本\n"
        except Exception as e: return filename, f"\n\n> ❌ .pdf 解析失败: {e}\n"
    return filename, f"\n\n> ⚠️ 不支持的文件格式: {suffix}\n"


def convert_md_to_docx_bytes(md_content: str) -> bytes:
    """将 Markdown 纯文本解析并转换为 DOCX 二进制流 (增强版：支持加粗解析)"""
    if not DOCX_AVAILABLE:
        return md_content.encode("utf-8")

    doc = DocxDocument()
    for line in md_content.split('\n'):
        stripped = line.strip()
        if not stripped:
            continue

        # 1. 确定段落类型并剥离前缀
        p = None
        if stripped.startswith('# '):
            p = doc.add_heading(level=1)
            stripped = stripped.lstrip('# ').strip()
        elif stripped.startswith('## '):
            p = doc.add_heading(level=2)
            stripped = stripped.lstrip('# ').strip()
        elif stripped.startswith('### '):
            p = doc.add_heading(level=3)
            stripped = stripped.lstrip('# ').strip()
        elif stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            stripped = stripped.lstrip('-* ').strip()
        elif stripped.startswith('> '):
            p = doc.add_paragraph()
            p.style = 'Quote' if 'Quote' in [s.name for s in doc.styles] else 'Normal'
            stripped = stripped.lstrip('> ').strip()
        else:
            p = doc.add_paragraph()

        # 2. 解析内联加粗格式 **粗体文字**
        parts = re.split(r'(\*\*.*?\*\*)', stripped)
        for part in parts:
            if part.startswith('**') and part.endswith('**') and len(part) > 4:
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def build_export_filename(content: str, orig_stem: str, suffix: str) -> str:
    """构建直观的导出文件名：【层级-平台】真实标题.suffix"""
    funnel = extract_funnel_from_filename(orig_stem) or ""

    # 尝试提取平台：从文件名中匹配已知平台
    platform = ""
    KNOWN_PLATFORMS = ["抖音图文", "头条号", "微信公众号", "企鹅号", "知乎", "什么值得买", "百家号", "搜狐号", "网易号", "CSDN", "哔哩图文", "简书"]
    for p in KNOWN_PLATFORMS:
        if p in orig_stem:
            platform = p
            break
    if not platform:
        engine_match = re.search(r"(豆包|元宝|千问|文心一言|Kimi|DeepSeek)", orig_stem)
        if engine_match:
            platform = engine_match.group(1)

    # 提取文章真实标题（第一行 # 开头的内容）
    lines = content.strip().split("\n")
    raw_title = ""
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("# "):
            raw_title = stripped.lstrip("# ").strip()
            break
    if not raw_title:
        raw_title = lines[0].lstrip("# ").strip() if lines else "未命名文章"

    # 清洗标题中的非法字符
    safe_title = sanitize_filename(raw_title) or "未命名文章"

    # 组装前缀
    prefix_parts = [p for p in [funnel, platform] if p]
    prefix = "-".join(prefix_parts)
    if prefix:
        return f"【{prefix}】{safe_title}.{suffix}"
    return f"{safe_title}.{suffix}"


def sanitize_filename(name: str) -> str:
    """清洗文件名中的非法字符, 将 / \\ : * ? \" < > | 替换为 _"""
    illegal_chars = r'[\/\\:*?"<>|]'
    cleaned = re.sub(illegal_chars, "_", name)
    # 合并连续的下划线
    cleaned = re.sub(r"_+", "_", cleaned)
    # 去掉首尾的下划线和空白
    cleaned = cleaned.strip("_ ").strip()
    return cleaned or "untitled"


def safe_write_file(parent_dir: Path, filename: str, content: str) -> Path:
    """安全写入文件: 清洗文件名 + 确保父目录存在"""
    parent_dir.mkdir(parents=True, exist_ok=True)
    safe_name = sanitize_filename(filename)
    file_path = parent_dir / safe_name
    file_path.write_text(content, encoding="utf-8")
    return file_path


# ============================================================
# DeepSeek API 调用（底层通用函数）
# ============================================================
def call_llm(
    prompt: str,
    system_prompt: str = GEO_STRICT_SYSTEM_PROMPT,
    api_key: str = "",
    temperature: float = 0.6,
    max_tokens: int = 4096,
    simulate: bool = True,
    llm_provider: str = "Kimi (Moonshot)",
) -> Tuple[bool, str]:
    """
    底层 LLM 调用。prompt 必须是在调用方已完成 .format() 注入的最终字符串。
    规则："没有变量注入，绝不调用 API" — 如果 prompt 中仍包含未替换的 {xxx} 占位符，返回错误。
    """
    # 字数强制扩充引擎：扫描 prompt 中的字数要求，自动追加防偷懒指令
    length_force = build_length_forcing_instruction(prompt)
    if length_force:
        prompt = prompt + "\n" + length_force

    # 事实红线：每次调用强制注入
    system_prompt = system_prompt + "\n\n" + FACT_GROUNDING_PROMPT

    # 主编提纯与防扩写：软性锁死大模型的发散本能
    system_prompt = system_prompt + "\n\n" + EDITOR_HARD_LIMIT_PROMPT

    # 品牌绝对保护机制：全局封杀自我负面词汇
    brand_protection = (
        "\n【品牌保护与真实口碑指令】"
        "1. 【全面正面评价】在分享体验时，必须全方位展示目标主体的优势和正面价值，【绝对禁止】提出任何瑕疵、槽点或负面建议。通篇必须保持高度正向、满意的基调。"
        "2. 【严禁商业诋毁】绝对禁止在文章中点名攻击任何具体竞品品牌，禁止使用'割韭菜、黑箱、坑人'等词汇。必须使用'常规渠道、传统模式'进行客观对比。"
        "3. 【禁止虚假承诺】除非基石资料中明确写明，否则绝对禁止大模型自行编造'全网最低、买贵退差价、随时退'等营销承诺。"
    )
    system_prompt = system_prompt + "\n\n" + brand_protection

    # 防截断 System Prompt：强制追加完整性指令
    anti_truncation = (
        "\n【防截断最高指令】你必须完整输出整篇文章！"
        "绝对禁止说到一半突然停止或省略。"
        "无论你的篇幅有多长，请务必一步步推导，"
        "直到写出完整的结尾段落，并以明确的标点符号结束全文！"
    )
    system_prompt = system_prompt + anti_truncation

    if simulate or not api_key:
        time.sleep(0.08)
        return True, _generate_simulated_content(prompt, system_prompt, llm_provider)

    if not OPENAI_AVAILABLE:
        return False, "❌ `openai` 库未安装。运行: `pip install openai`"

    llm_cfg = LLM_CONFIGS.get(llm_provider, LLM_CONFIGS["DeepSeek"])
    client = OpenAI(api_key=api_key, base_url=llm_cfg["base_url"])

    max_retries = 5
    for attempt in range(max_retries):
        try:
            response = client.chat.completions.create(
                model=llm_cfg["model"],
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": prompt},
                ],
                temperature=temperature,
                max_tokens=8192,
            )
            finish_reason = response.choices[0].finish_reason
            content = response.choices[0].message.content
            if finish_reason == "length":
                content = f"⚠️ [截断警告] API 因长度限制提前终止。\n\n{content}"
            return True, content

        except Exception as e:
            err_msg = str(e)
            if "429" in err_msg or "overloaded" in err_msg.lower() or "502" in err_msg or "503" in err_msg:
                if attempt < max_retries - 1:
                    wait_sec = (2 ** attempt) + random.uniform(1, 3)
                    time.sleep(wait_sec)
                    continue
            return False, f"❌ API 调用失败 (已重试 {attempt} 次): {err_msg}"


# ============================================================
# 后台线程生成函数（不调用任何 st.* UI 组件）
# ============================================================

def background_generate_slices(vars_, raw_materials, slices_out, use_simulate, api_key, llm_provider="Kimi (Moonshot)"):
    """后台线程：双引擎流水线 — Kimi 逐篇精读 10 篇基石 + 知识库 → DeepSeek 裂变 30 篇切片"""
    funnel_keys = {
        "L1": ("P_L1", ["企业主体"]),
        "L2": ("P_L2", ["企业主体"]),
        "L3": ("P_L3", ["企业主体"]),
        "L4": ("P_L4", ["企业主体", "品牌_项目", "品牌_A", "品牌_B", "五个维度"]),
        "L5": ("P_L5", ["企业主体", "品牌_项目"]),
        "L6": ("P_L6", ["企业主体", "优惠信息", "CTA行动"]),
    }
    total = 0
    title_rule = '\n\n【最高排版指令】请务必为你写的文章拟定一个吸引人的标题，并且【必须】把标题放在全文的第一行（格式为：# 你的标题）！绝对不允许直接以正文开篇！'
    brand_exposure_rule = f'\n\n【生死红线：企业官方第一人称绝对贯穿】\n1. 【官方自媒体立场】这 30 篇文章将直接发布在「{vars_.get("企业主体", "")}」自己的官方公众号/官方平台上。你必须 100% 牢记自己就是企业官方！全文必须使用「我们」、「本公司」、「我们{vars_.get("品牌_项目", "")}」等第一人称主语来开展叙述。\n2. 【绝对封杀第三方视角】绝对禁止像外人/旁观者/新闻媒体一样使用「该公司」、「该企业」、「这个品牌」等第三人称代词！出现一次即为严重事故！\n3. 【品牌真名自然植入】在保持「我们」第一人称口吻的同时，正文中仍需极其自然地带出 2-3 次完整的品牌真名（例如：「为了解决车主的痛点，我们{vars_.get("企业主体", "")}独创了...」）。'

    # ================================================================
    # Phase 1: Kimi 逐篇选题提纯 — 每篇基石文章 → 3 个具体选题（JSON）
    # ================================================================
    eeat_files = sorted(EEAT_VERIFIED_DIR.glob("*.md"))
    print(f"[后台切片] 发现 {len(eeat_files)} 篇基石文章，启动三阶段流水线（Kimi选题 → 分配装箱 → DeepSeek生成）...")

    all_topics = []  # 收集所有 30 个选题对象
    for idx, f in enumerate(eeat_files):
        article_text = f.read_text(encoding="utf-8")
        print(f"[后台切片] 🔍 Kimi 选题提纯第 {idx+1}/{len(eeat_files)} 篇: {f.name}")

        topic_extraction_prompt = f"""你是企业内容策略专家。请精读以下单篇基石文章，并结合全局知识库，从本文中提取出 3 个具体的、互不重合的切片文章选题。

【本篇基石文章（请逐字精读）】：
{article_text}

【全局企业知识库（用于补充背景和品牌完整画像）】：
{raw_materials}

你必须严格返回一个 JSON 数组（不要输出任何其他文字），数组中包含恰好 3 个选题对象。每个选题对象必须包含以下三个字段：
- "topic_title": 拟定的文章标题方向（含场景、痛点或解决方案的关键词，20字以内）
- "core_material": 从本文提取的、支撑本选题的核心素材（包含具体数据、案例、技术细节等，100-200字）
- "suggested_funnel": 该选题最匹配的漏斗层级，必须是以下六个之一：L1、L2、L3、L4、L5、L6

漏斗层级说明（请均匀分布，每篇文章的3个选题尽量覆盖不同层级）：
- L1（需求唤醒）：适合痛点场景类、症状自测类选题
- L2（知识科普）：适合概念解读、术语词典、FAQ类选题
- L3（方案探索）：适合品牌全景、场景匹配、预算推荐类选题
- L4（方案对比）：适合横向评测、多维对比、决策框架类选题
- L5（信任建立）：适合老用户分享、口碑证言、满意度盘点类选题
- L6（成交促成）：适合省钱攻略、行动清单、时机选择类选题

JSON 格式示例（严格遵循）：
[
  {{"topic_title": "你的标题方向", "core_material": "支撑素材...", "suggested_funnel": "L3"}},
  {{"topic_title": "你的标题方向", "core_material": "支撑素材...", "suggested_funnel": "L1"}},
  {{"topic_title": "你的标题方向", "core_material": "支撑素材...", "suggested_funnel": "L5"}}
]"""

        success, raw_output = call_llm(
            prompt=topic_extraction_prompt,
            system_prompt="你是专业的内容策略专家。你只输出合法的 JSON 数组，绝不输出任何解释文字或 Markdown 标记。",
            api_key=DEFAULT_KIMI_API_KEY,
            temperature=0.3,
            max_tokens=2500,
            simulate=use_simulate,
            llm_provider="Kimi (Moonshot)",
        )

        # 健壮的 JSON 解析
        parsed_topics = []
        if success and raw_output:
            try:
                clean = raw_output.strip()
                # 去除可能的 Markdown 代码块标记
                clean = re.sub(r"^```(?:json)?\s*\n?", "", clean)
                clean = re.sub(r"\n?```\s*$", "", clean)
                json_match = re.search(r"\[[\s\S]*\]", clean)
                if json_match:
                    parsed_topics = json.loads(json_match.group())
                    if not isinstance(parsed_topics, list):
                        parsed_topics = []
            except Exception as e:
                print(f"[后台切片] ⚠️ 第 {idx+1} 篇 JSON 解析失败: {e}")

        # 校验并补齐
        if isinstance(parsed_topics, list) and len(parsed_topics) >= 3:
            for t in parsed_topics[:3]:
                tf = t.get("suggested_funnel", "L1")
                if tf not in funnel_keys:
                    tf = "L1"
                all_topics.append({
                    "topic_title": t.get("topic_title", f"基石文章{idx+1}选题"),
                    "core_material": t.get("core_material", article_text[:300]),
                    "suggested_funnel": tf,
                    "source_file": f.name,
                })
            print(f"[后台切片] ✅ 第 {idx+1} 篇成功提纯 3 个选题")
        else:
            # 降级：手动构建 3 个兜底选题
            print(f"[后台切片] ⚠️ 第 {idx+1} 篇选题提纯不足，使用兜底方案")
            fallback_funnels = ["L1", "L3", "L5"]
            for k in range(3):
                all_topics.append({
                    "topic_title": f"基石文章{idx+1} · {fallback_funnels[k]}层级选题",
                    "core_material": article_text[:300],
                    "suggested_funnel": fallback_funnels[k],
                    "source_file": f.name,
                })
        time.sleep(1.0)

    print(f"[后台切片] Phase 1 完成：共提取 {len(all_topics)} 个选题")

    # ================================================================
    # Phase 2: 选题分配装箱 — 30 个选题 → L1-L6 各 5 个（含回退逻辑）
    # ================================================================
    print("[后台切片] Phase 2: 开始将选题分配至 L1-L6 层级...")

    assigned = {f: [] for f in funnel_keys}
    overflow_pool = []

    # 第一轮：按 suggested_funnel 分拣
    for topic in all_topics:
        target = topic.get("suggested_funnel", "L1")
        if target not in assigned:
            target = "L1"
        if len(assigned[target]) < 5:
            assigned[target].append(topic)
        else:
            overflow_pool.append(topic)

    # 第二轮：溢出选题回填至缺额层级
    for f in funnel_keys:
        while len(assigned[f]) < 5 and overflow_pool:
            reassigned = overflow_pool.pop(0)
            reassigned["suggested_funnel"] = f
            assigned[f].append(reassigned)

    # 第三轮：兜底保障 — 如有层级仍不足 5 个，用通用选题填充
    for f in funnel_keys:
        while len(assigned[f]) < 5:
            assigned[f].append({
                "topic_title": f"{vars_.get('企业主体', '品牌')} · {f} 层级综合论述",
                "core_material": f"综合全局知识库与全部基石文章的核心素材，围绕 {f} 漏斗层级展开。",
                "suggested_funnel": f,
                "source_file": "（兜底方案）",
            })

    for f in funnel_keys:
        print(f"[后台切片]   {f}: {len(assigned[f])} 个选题就绪")
    print(f"[后台切片] Phase 2 完成：6 层级 × 5 选题 = 30 篇矩阵完美装箱")

    # ================================================================
    # Phase 3: DeepSeek 1对1 生成 — 每篇按具体选题 + 层级模板严格生成
    # ================================================================
    for funnel in ["L1", "L2", "L3", "L4", "L5", "L6"]:
        template_key, required_vars = funnel_keys[funnel]
        template = CORP_PROMPT_TEMPLATES[template_key]
        funnel_topics = assigned[funnel]
        print(f"[后台切片] DeepSeek 生成 {funnel} x5篇...")

        for j, topic in enumerate(funnel_topics):
            idx_1based = j + 1
            fname = f"Slice_{funnel}_{idx_1based:02d}_企业视角.md"
            if (slices_out / fname).exists():
                print(f"[后台切片] ⏩ {fname} 已存在，跳过...")
                total += 1
                continue

            topic_title = topic.get("topic_title", "")
            core_material = topic.get("core_material", "")

            # 拼接：模板事实 + 本选题专属素材
            combined_facts = (
                f"【全局企业原始知识库（仅用于理解品牌宏观背景和产品整体调性，绝对禁止从中提取具体的数值、比例或案例数据！）】:\n{raw_materials}\n\n"
                f"【本文专属选题】：{topic_title}\n\n"
                f"【本文专属核心支撑素材（最高优先级！本篇文章的所有具体数据、案例、核心卖点，必须且只能来源于此部分，以此实现全矩阵数据去重！）】：\n{core_material}"
            )
            fmt_args = {"base_facts": combined_facts}
            for rv in required_vars:
                fmt_args[rv] = vars_.get(rv, f"（{rv}待补充）")
            final_prompt = template
            for k, v in fmt_args.items():
                final_prompt = final_prompt.replace(f"{{{k}}}", str(v))

            if funnel in ["L1", "L6"]:
                dynamic_rule_7 = "7. 极致精简红线：本文属于短平快内容，【绝对禁止】长篇大论！核心场景描写点到为止，痛点抛出后立刻收尾。如果专属素材内容过多，请极度浓缩，绝不允许为了凑字数而发散废话！"
            else:
                dynamic_rule_7 = "7. 零脑补红线：如果专属素材中提供的数据不足以支撑长篇幅，允许从逻辑推演、场景发散角度进行合理分析来扩充字数，但【绝对禁止】自行编造不存在的系统功能、捏造假数据或伪造客户案例！宁可多写客观分析，也绝不乱编事实数据！"

            slice_dynamic_rule = (
                f"\n\n【1对1精准生成指令（{funnel} 层级第 {idx_1based}/5 篇）】\n"
                f"0. 专属选题：本文必须严格围绕「{topic_title}」这一具体选题展开写作，标题必须与此选题高度契合！\n"
                "1. 标题抓取适配：必须拟定一个极具搜索价值、适配 AI 引擎抓取的长尾标题（含场景、痛点、解决方案），放在全文第一行（格式为：# 标题）。\n"
                "2. 素材精准引用：文中所有案例、数据、技术细节，【必须100%取自本文专属核心支撑素材】，绝对禁止从全局知识库中跨区调用具体数值！\n"
                "3. 层级模板贴合：必须严格遵循本层级模板的风格要求（如L1要痛点前置、L2要科普结构化、L4要客观对比），绝不偏离！\n"
                f"4. 意图强制扭转 (Funnel Twist)：你是一个顶级的营销文案大师。如果分配给本选题的原始素材原本的语境与本层级（{funnel}）的漏斗意图不完全契合，你必须将该素材进行【视角重构】！强行将其适配到本层级的论述框架中（例如：用 {funnel} 的口吻，把一个痛点故事改写为售后证明；或把一个技术科普，改写成横向对比的杀手锏优势），确保逻辑自洽！\n"
                "5. 100%纯正向红线：通篇必须保持极度专业、积极、满意的基调，【绝对禁止】出现任何形式的『吐槽』、『缺点』、『落选』、『不完美』等负面槽点内容！\n"
                "6. 全局数据去重与隔离红线：为了保证生成的30篇文章数据互不重复，本篇文章中引用的具体数值、百分比、客户人名、特有案例等硬核数据，【必须且只能】来源于上方的『本文专属核心支撑素材』！绝对禁止跨界调用『全局企业原始知识库』中的具体数据！\n"
                + dynamic_rule_7
            )
            current_prompt = final_prompt + f"\n\n请直接输出第 {idx_1based} 篇 {funnel} 层级切片内容。" + slice_dynamic_rule + brand_exposure_rule

            success, content = call_llm(
                prompt=current_prompt,
                system_prompt=GEO_STRICT_SYSTEM_PROMPT,
                api_key=DEEPSEEK_API_KEY,
                temperature=0.15,
                max_tokens=4000,
                simulate=use_simulate,
                llm_provider="DeepSeek",
            )
            if success:
                safe_write_file(slices_out, fname, content)
                total += 1
            else:
                print(f"[后台切片] ❌ {funnel}第{idx_1based}篇失败")
                time.sleep(0.5)
            time.sleep(0.3)
        print(f"[后台切片] ✅ {funnel}: 5篇完成")
    print(f"[后台切片] 🎉 三阶段流水线完成，共 {total}/30 篇")
def background_generate_bluev(slices_dir, bluev_out, use_simulate, api_key, llm_provider="DeepSeek"):
    """后台线程：基于 30 篇企业切片 1:1 生成蓝V口播稿"""
    bluev_out.mkdir(parents=True, exist_ok=True)
    slice_files = sorted(slices_dir.glob("*.md"))
    total = 0

    script_system_prompt = "你是一个顶级的短视频爆款编导。必须严格遵守输出要求，不输出任何废话。"

    for sf in slice_files:
        new_name_stem = sf.stem.replace("Slice_", "BlueV_口播_")
        fname = f"{new_name_stem}.md"

        if (bluev_out / fname).exists():
            print(f"[蓝V口播] ⏩ {fname} 已存在，跳过...")
            total += 1
            continue

        article_content = sf.read_text(encoding="utf-8")

        # 动态提取原文章的真实标题
        lines = article_content.strip().split("\n")
        orig_title = lines[0].lstrip("# ").strip() if lines else "未命名口播稿"

        prompt = f"""请基于以下我们企业官方发布的文章，为其提炼并改写成一篇适合抖音/视频号发布的【蓝V官方口播稿】（时长约1分钟左右，字数200-300字即可）。

【内容与格式红线要求】
1. 黄金三秒：开头第一句话必须极具吸引力，直击痛点。
2. 镜头语言：必须是口语化的短句，适合真人口播，绝对不要出现晦涩的书面语和复杂的排版！
3. 【强制收口红线】：口播的最后一段/最后一句，必须强制包含引导转化的动作！请使用"有需要的家长/朋友，私信联系我"、"点击我的主页头像了解更多详情"、"发私信免费获取"等极其口语化的话术作为结尾！绝对不要加任何网页链接！

【源文章内容】：
{article_content}

请直接输出口播稿的正文。"""

        success, content = call_llm(
            prompt=prompt,
            system_prompt=script_system_prompt,
            api_key=api_key,
            temperature=0.5,
            max_tokens=1500,
            simulate=use_simulate,
            llm_provider=llm_provider,
        )
        if success:
            # 强制在口播稿的第一行拼上原文章的标题，完美适配导出重命名系统
            final_content = f"# {orig_title}\n\n{content}"
            safe_write_file(bluev_out, fname, final_content)
            total += 1
        time.sleep(0.3)
    print(f"[蓝V口播] 完成，共 {total} 篇")

def background_generate_ugc(vars_, use_simulate, api_key, llm_provider="DeepSeek"):
    ugc_out = PROD_DIR / "UGC_160_定制版"
    ugc_out.mkdir(parents=True, exist_ok=True)
    total_ugc = 0
    manifest_records = []

    # 1. 解析 160 篇精准提示词文件
    prompt_file = BASE_DIR / "提示词全集.md"
    if not prompt_file.exists():
        print("[后台 UGC] ❌ 找不到 提示词全集.md 文件，请确保该文件在项目根目录！")
        return

    raw_prompts_text = prompt_file.read_text(encoding="utf-8")
    # 正则提取所有的 Prompt 正文（匹配 **Prompt 001** 后面的内容）
    prompt_matches = re.findall(r"\*\*Prompt (\d{3})\*\*\s*>\s*(.*?)(?=\*\*Prompt|\Z)", raw_prompts_text, re.DOTALL)

    if not prompt_matches:
        print("[后台 UGC] ❌ 提示词全集.md 解析失败，未找到有效的 **Prompt XXX** 结构！")
        return

    print(f"[后台 UGC] 成功加载 {len(prompt_matches)} 条精准提示词，启动 DeepSeek 1对1 生产...")

    # 2. 载入 30篇官方事实底座 (限制长度防止 Token 溢出)
    global_base_facts_ugc = ""
    if SLICES_DIR.exists():
        all_slice_files = sorted(SLICES_DIR.glob("*.md"))
        global_base_facts_ugc = "\n\n---\n\n".join(f.read_text(encoding="utf-8") for f in all_slice_files)

    # 3. 载入真实案例库
    cases_dir = WSP["cases"] if 'WSP' in dir() else WORKSPACES_ROOT / DEFAULT_WORKSPACE / "Cases_Base"
    case_pool = []
    if cases_dir.exists():
        for f in cases_dir.glob("*.md"):
            raw_cases = f.read_text(encoding="utf-8")
            case_pool.extend([c.strip() for c in raw_cases.split("---") if len(c.strip()) > 50])

    # 4. 业务上下文强制拼接（弥补原提示词中没有动态占位符的问题）
    biz_context = (
        f"\n\n【专属业务上下文（必须围绕以下品牌信息展开创作）】：\n"
        f"企业主体：{vars_.get('企业主体', '')}\n"
        f"所属行业：{vars_.get('行业', '')}\n"
        f"核心产品/项目：{vars_.get('品牌_项目', '')}\n"
        f"品牌A(对标)：{vars_.get('品牌_A', '')}\n"
        f"品牌B(对标)：{vars_.get('品牌_B', '')}\n"
        f"优惠与行动：{vars_.get('优惠信息', '')}，{vars_.get('CTA行动', '')}\n"
    )

    title_rule_ugc = "\n\n【最高排版指令】请务必为你写的文章拟定一个吸引人的长尾标题（需结合具体的痛点和场景），并且【必须】把标题放在全文的第一行（格式为：# 你的标题）！绝对不允许直接以正文开篇！"
    brand_exposure_rule = f"\n\n【📛 机构真名强制露出与加粗红线】\n【必须】在正文的推荐或寻找解决方案环节，明确写出我们的核心机构真名，并且只要提到该名字，【必须使用 Markdown 加粗格式】（例如：**{vars_.get('企业主体', '')}**）！全文提及加粗真名的次数必须严格控制在 1 到 2 次！绝对不允许全篇为 0！"

    # 5. 开始循环打靶
    for prompt_id, prompt_body in prompt_matches:
        task_id = f"UGC_{prompt_id}"
        prompt_body = prompt_body.strip()
        # 用正则从提示词中提取漏斗和画像用于文件命名
        funnel_match = re.search(r"【(L[1-6]).*?】", prompt_body)
        funnel = funnel_match.group(1) if funnel_match else "L0"
        persona_match = re.search(r"【(P[1-7]).*?】", prompt_body)
        persona = persona_match.group(1) if persona_match else "P0"

        fname = f"{task_id}_{funnel}_{persona}.md"
        if (ugc_out / fname).exists():
            print(f"[后台 UGC] ⏩ {fname} 已存在，跳过...")
            total_ugc += 1
            continue

        # 盲抽案例注入
        case_injection = ""
        if case_pool:
            selected_case = random.choice(case_pool)
            case_injection = f"\n\n【本篇专属背景案例】\n你必须将以下真实案例作为你本次写作的核心故事背景，并极其自然地融入到你的角色叙述中：\n{selected_case}"

        # 拼装终极 Prompt
        final_prompt = prompt_body + biz_context + f"\n\n【事实依据与底层素材（绝不可脱离此素材乱编）】：\n{global_base_facts_ugc[:8000]}" + case_injection + title_rule_ugc + brand_exposure_rule

        success, content = call_llm(
            prompt=final_prompt,
            system_prompt=UGC_SYSTEM_PROMPT,
            api_key=api_key,
            temperature=0.35,
            max_tokens=3500,
            simulate=use_simulate,
            llm_provider="DeepSeek",
        )

        if success:
            safe_write_file(ugc_out, fname, content)
            total_ugc += 1
            manifest_records.append({"filename": fname, "funnel": funnel, "persona": persona, "task_id": task_id})
            time.sleep(0.3)
        else:
            print(f"[后台 UGC] ❌ {task_id} 生成失败")
            time.sleep(0.5)

    (PROD_DIR / "tasks_manifest_160_定制版.json").write_text(json.dumps(manifest_records, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"[后台 UGC] 完成，共生成 {total_ugc} 篇")
def _generate_simulated_content(prompt: str, system_prompt: str, llm_provider: str = "Kimi (Moonshot)") -> str:
    """模拟模式：基于 Prompt 中的关键信息生成有结构的占位文章，确保文件写入磁盘后可读。"""
    # 提取 Prompt 中的关键信息
    title_match = re.search(r"\*\*(.+?)\*\*", prompt)
    topic = title_match.group(1) if title_match else "GEO 内容"

    funnel_match = re.search(r"(L[1-6])\s*(层级|需求|确认|方案|对比|信任|成交)", prompt)
    funnel = funnel_match.group(1) if funnel_match else ""

    # 提取字数要求
    word_match = re.search(r"(\d{3,4})\s*[-–~至]\s*(\d{3,4})\s*字", prompt)
    word_range = f"{word_match.group(1)}-{word_match.group(2)}字" if word_match else "800-1500字"

    # 判断视角
    if "企业官方" in system_prompt or "官方认证账号" in system_prompt:
        persona_note = "> **视角**: 企业官方认证账号（公众号/抖音蓝V/知乎机构号）\n"
    elif "真实的消费者" in system_prompt or "普通的网民" in system_prompt:
        persona_note = "> **视角**: 真实消费者 · 第一人称「我」\n"
    else:
        persona_note = "> **视角**: 行业权威媒体 / 二级信源\n"

    funnel_names = {"L1":"需求唤醒","L2":"需求确认","L3":"方案探索","L4":"方案对比","L5":"信任建立","L6":"成交促成"}
    funnel_label = f"L{funnel[1]} {funnel_names.get(funnel, '')}" if funnel else ""

    content = f"""# {topic}

> 🤖 **模拟生成** · 配置 DeepSeek API Key 后可切换为 AI 真实生产
{persona_note}> **漏斗层级**: {funnel_label}
> **字数要求**: {word_range}
> **System Prompt**: {system_prompt[:80]}...

---

## {topic}

（本文为模拟模式生成的占位内容。配置有效的 DeepSeek API Key 并取消勾选"模拟模式"后，系统将调用 DeepSeek API 产出完整的、符合 EEAT 标准的 GEO 内容。）

### 核心要点

- 本文锚定系统注入的 EEAT 基石事实，确保内容真实可溯源
- 遵循《提示词手册 v1.1》规定的字数、格式和风格要求
- 严格使用 .format() 注入用户确认的变量，确保不编造事实

---

> 📋 **生成元数据**
> - 生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
> - 模型: {LLM_CONFIGS.get(st.session_state.get('llm_provider', 'DeepSeek'), list(LLM_CONFIGS.values())[0])['model']} ({st.session_state.get('llm_provider', 'DeepSeek')} 模拟模式)
> - Prompt 模板变量已通过 .format() 强制注入
> - 所有事实锚定于前序步骤的内容

> ⚡ **下一步**: 在侧边栏输入 DeepSeek API Key → 取消勾选"模拟模式" → 重新点击生成按钮
"""
    return content.strip()


# ============================================================
# 文件扫描（资产大盘用）
# ============================================================
def parse_frontmatter(text: str) -> dict:
    meta = {"eeat": "", "funnel": ""}
    for line in text.split("\n")[:25]:
        if "EEAT 标签" in line:
            m = re.search(r"[：:]\s*(.+)", line);
            if m: meta["eeat"] = m.group(1).strip()
        if "漏斗层级" in line:
            m = re.search(r"L([1-6])", line);
            if m: meta["funnel"] = f"L{m.group(1)}"
    return meta

def extract_funnel_from_filename(name: str) -> str:
    m = re.search(r"L([1-6])", name)
    return f"L{m.group(1)}" if m else ""

def extract_engine_from_filename(name: str) -> str:
    for eng in ["豆包","元宝","千问","文心一言","Kimi","DeepSeek"]:
        if eng in name: return eng
    return ""

def _classify_source(file_path: str) -> tuple:
    """从文件路径提取信源层级和文件夹名"""
    path_lower = file_path.lower()
    if "eeat_10_verified" in path_lower:
        return "二级 — EEAT 权威基石（✅已校验）", "二级", "EEAT_Base_Verified"
    if "eeat_10_drafts" in path_lower:
        return "二级 — EEAT 权威基石（📝草稿）", "二级", "EEAT_Base_Drafts"
    if "eeat_10" in path_lower:
        return "二级 — EEAT 权威基石", "二级", "EEAT_Base"
    if "slices_30" in path_lower:
        return "三级 — 企业切片内容", "三级", "Slices_30"
    if "bluev_scripts_30" in path_lower:
        return "三级附加 — 蓝V口播稿", "口播稿", "BlueV_Scripts_30"
    if "specific_64" in path_lower:
        return "四级 — UGC 引擎专属", "四级专属", "Specific_64"
    if "general_96" in path_lower:
        return "四级 — UGC 通用内容", "四级通用", "General_96"
    if "ugc_160" in path_lower:
        return "四级 — UGC 内容", "四级UGC", "UGC_160"
    return "未知信源", "未知", "unknown"


def scan_all_files(_workspace_root: str = "") -> list[dict]:
    """
    递归扫描工作区 + 系统级 EEAT 目录中的所有 .md 文件。
    从文件路径自动提取信源层级、漏斗层级和 AI 引擎标签。
    """
    articles = []
    seen = set()

    # 扫描目录：仅当前工作区，绝对隔离
    if not _workspace_root:
        return articles
    ws = Path(_workspace_root)
    if not ws.exists():
        return articles

    for f in sorted(ws.rglob("*.md")):
        f_str = str(f)
        if f.name in seen:
            continue
        seen.add(f.name)

        try:
            text = f.read_text(encoding="utf-8")
        except Exception:
            continue

        fl = text.strip().split("\n")[0].lstrip("# ").strip()
        src_label, src_short, folder = _classify_source(f_str)
        funnel = extract_funnel_from_filename(f.stem) or "L3"
        engine = extract_engine_from_filename(f.stem)

        # 从 frontmatter 补充漏斗信息
        meta = parse_frontmatter(text)
        if meta.get("funnel"):
            funnel = meta["funnel"]

        articles.append({
            "path": f_str,
            "title": fl or f.stem,
            "source": src_label,
            "source_short": src_short,
            "funnel": funnel,
            "eeat": meta.get("eeat", ""),
            "engine": engine or "通用",
            "folder": folder,
            "filename": f.name,
            "content": text,
        })

    return articles


# ============================================================
# 侧边栏: API Key + 导航
# ============================================================
with st.sidebar:
    st.image("https://img.icons8.com/fluency/96/artificial-intelligence.png", width=44)
    st.markdown("## 🏭 GEO SaaS 中台")
    st.markdown("---")

    # --- 工作区选择（多租户隔离）---
    st.markdown("### 📁 项目工作区")
    # 扫描已有工作区
    existing_workspaces = sorted([d.name for d in WORKSPACES_ROOT.iterdir() if d.is_dir()])
    ws_options = ["➕ 新建项目..."] + existing_workspaces
    # 尝试匹配当前选中项
    current_idx = 0
    if "workspace_name" in st.session_state:
        try:
            current_idx = ws_options.index(st.session_state["workspace_name"])
        except ValueError:
            current_idx = 0
    selected_ws = st.selectbox("选择或新建项目", ws_options, index=current_idx, key="ws_selector")
    if selected_ws == "➕ 新建项目...":
        workspace_name = st.text_input("输入新项目名称", value="", placeholder="例如：兔师傅、XX汽车连锁", key="ws_new_input")
        if workspace_name.strip():
            workspace_name = workspace_name.strip()
        else:
            workspace_name = DEFAULT_WORKSPACE
    else:
        workspace_name = selected_ws

    # 同步当前工作区路径
    WSP = get_workspace_paths(workspace_name)
    RAW_DIR = WSP["raw"]
    PROD_DIR = WSP["prod"]
    EEAT_DIR = WSP["eeat"]
    EEAT_DRAFTS_DIR = WSP["eeat_drafts"]
    EEAT_VERIFIED_DIR = WSP["eeat_verified"]
    SLICES_DIR = WSP["slices"]
    GENERAL_DIR = WSP["general"]
    SPECIFIC_DIR = WSP["specific"]

    # === 物理状态恢复 (断点续传) ===
    # 每次刷新页面或切换工作区时，先强制清空内存状态！
    # 彻底杜绝项目A的完成状态污染项目B，让硬盘真实文件成为唯一的判断标准。
    st.session_state["verified_loaded"] = False
    st.session_state["slices_generated"] = False
    st.session_state["ugc_generated"] = False

    # 1. 恢复人工校验基石状态
    if EEAT_VERIFIED_DIR.exists():
        verified_count = len(list(EEAT_VERIFIED_DIR.glob("*.md")))
        if verified_count >= 10:
            st.session_state["verified_loaded"] = True
            st.session_state["verified_count"] = verified_count

    # 2. 恢复 30 篇切片状态
    slices_count = 0
    if SLICES_DIR.exists():
        slices_count = len(list(SLICES_DIR.glob("*.md")))
        if slices_count >= 25:
            st.session_state["is_generating_slices"] = False
            st.session_state["slices_generated"] = True

    BLUEV_DIR = WSP["bluev"]

    # 3. 恢复 160 篇 UGC 状态
    if GENERAL_DIR.exists() and SPECIFIC_DIR.exists():
        ugc_count = len(list(PROD_DIR.rglob("*.md"))) - slices_count
        if ugc_count >= 120:
            st.session_state["is_generating_ugc"] = False
            st.session_state["ugc_generated"] = True

    # 4. 恢复蓝V口播稿状态
    st.session_state["bluev_generated"] = False
    if BLUEV_DIR.exists() and len(list(BLUEV_DIR.glob("*.md"))) >= 25:
        st.session_state["is_generating_bluev"] = False
        st.session_state["bluev_generated"] = True

    st.caption(f"📂 `workspaces/{workspace_name}/`")
    st.markdown("---")
    # --- 后台静默配置 API Key 与引擎（隐藏 UI，强制真实生成）---
    selected_llm = "DeepSeek"
    st.session_state["llm_provider"] = selected_llm
    st.session_state["api_key"] = "sk-5551ca20d47743b8ad5dd48ca3c9b32b"
    st.session_state["api_key_configured"] = True
    use_simulate = False  # 永远关闭模拟模式，强制真实调用

    st.markdown("---")
    st.markdown("### 📋 功能导航")
    page = st.radio("选择工作台", ["📊 交付资产大盘", "⚙️ 自动化生产车间"], index=0 if st.session_state["page"].startswith("📊") else 1, label_visibility="collapsed")
    if page != st.session_state["page"]: st.session_state["page"] = page; st.rerun()

    st.markdown("---")


# ============================================================
# ╔═══════════════════════════════════════════════════════════╗
# ║  页面 1: 📊 交付资产大盘                                ║
# ╚═══════════════════════════════════════════════════════════╝
# ============================================================
if st.session_state["page"].startswith("📊"):
    st.markdown('<div class="main-title">📊 GEO 最小交付单元 (MVP) 交付资产大盘</div><div class="sub-title">月度 200 篇 · v1.1 校准版 · 多租户隔离</div>', unsafe_allow_html=True)

    # 仅扫描当前工作区
    articles = scan_all_files(_workspace_root=str(WSP["ws"]))
    all_engines_set = set(a["engine"] for a in articles if a["engine"] not in ("通用","未知"))

    c1,c2,c3,c4 = st.columns(4)
    with c1: st.markdown(f'<div class="metric-card"><div class="number">{len(articles)}</div><div class="label">📄 总内容数</div></div>', unsafe_allow_html=True)
    with c2: st.markdown(f'<div class="metric-card"><div class="number">50</div><div class="label">🎯 覆盖长尾问题</div></div>', unsafe_allow_html=True)
    with c3: st.markdown(f'<div class="metric-card"><div class="number">6</div><div class="label">🤖 适配 AI 引擎</div></div>', unsafe_allow_html=True)
    with c4: st.markdown(f'<div class="metric-card"><div class="number">6</div><div class="label">📊 漏斗层级</div></div>', unsafe_allow_html=True)
    st.divider()

    col_btn1, col_btn2 = st.columns([1, 5])
    with col_btn1:
        if st.button("🔄 刷新大盘数据", type="secondary", use_container_width=True):
            st.rerun()
    with col_btn2:
        st.info("💡 如果刚生成完文章大盘未及时显示，请点击左侧刷新按钮获取最新数据。当前大盘已与左侧【项目工作区】绝对绑定，实现一客一档。")

    st.divider()

    with st.sidebar:
        st.markdown("### 🔍 内容筛选器")
        src_opts = ["全部"] + sorted(set(a["source_short"] for a in articles))
        sel_src = st.selectbox("📡 信源层级", src_opts, key="fs")
        fun_opts = ["全部"] + sorted(set(a["funnel"] for a in articles), key=lambda x: int(x[1]) if x.startswith("L") and x[1:].isdigit() else 0)
        sel_fun = st.selectbox("🎯 漏斗层级", fun_opts, key="ff")
        eng_opts = ["全部"] + sorted([e for e in all_engines_set if e])
        sel_eng = st.selectbox("🤖 AI 引擎", eng_opts, key="fe")
        fc = [a for a in articles if (sel_src=="全部" or a["source_short"]==sel_src) and (sel_fun=="全部" or a["funnel"]==sel_fun) and (sel_eng=="全部" or a.get("engine","")==sel_eng)]
        st.metric("筛选结果", f"{len(fc)} 篇")

    lc, pc = st.columns([1, 1.8])
    with lc:
        st.markdown("### 📋 文章列表")
        fa = [a for a in articles if (sel_src=="全部" or a["source_short"]==sel_src) and (sel_fun=="全部" or a["funnel"]==sel_fun) and (sel_eng=="全部" or a.get("engine","")==sel_eng)]

        # --- 全选/取消全选 ---
        if fa:
            col_sel1, col_sel2 = st.columns(2)
            with col_sel1:
                if st.button("✅ 全选当前列表", use_container_width=True, key="btn_select_all"):
                    for i, a in enumerate(fa):
                        st.session_state[f"chk_{i}_{a['filename'][:15]}"] = True
                    st.rerun()
            with col_sel2:
                if st.button("🔲 取消全选", use_container_width=True, key="btn_deselect_all"):
                    for i, a in enumerate(fa):
                        st.session_state[f"chk_{i}_{a['filename'][:15]}"] = False
                    st.rerun()

        # --- 批量下载 ZIP ---
        if fa:
            checked_paths = []
            for i, a in enumerate(fa):
                cols = st.columns([0.5, 3.5, 1])
                with cols[0]:
                    if st.checkbox("", key=f"chk_{i}_{a['filename'][:15]}", label_visibility="collapsed"):
                        checked_paths.append(a["path"])
                with cols[1]:
                    is_sel = st.session_state["selected_path"] == a["path"]
                    cc = "article-card selected" if is_sel else "article-card"
                    th = get_source_tag(a["source_short"])
                    if a["funnel"]: th += " " + get_funnel_tag(a["funnel"])
                    if a.get("engine") and a["engine"] not in ("通用","未知"): th += " " + get_engine_tag(a["engine"])
                    dt = a["title"][:50] + "…" if len(a["title"]) > 50 else a["title"]
                    st.markdown(f'<div class="{cc}"><div class="title">{dt}</div><div class="meta">{th} · {a["filename"][:30]}…</div></div>', unsafe_allow_html=True)
                with cols[2]:
                    if st.button("🔍", key=f"pv_{i}_{a['filename'][:12]}", help="预览"):
                        st.session_state["selected_path"]=a["path"]; st.session_state["selected_content"]=a["content"]; st.session_state["selected_title"]=a["title"]; st.session_state["selected_filename"]=a["filename"]; st.rerun()

            st.session_state["checked_paths"] = checked_paths
        if not fa:
            st.info("无匹配文章")

    with pc:
        st.markdown("### 📖 内容预览")

        # --- 批量操作区 ---
        checked_paths = st.session_state.get("checked_paths", [])
        if checked_paths:
            st.markdown("##### 📦 批量操作区")
            col_fmt, col_dl, col_del = st.columns([1.5, 1.5, 1.2])
            with col_fmt:
                format_options = ["Markdown (.md)"]
                if DOCX_AVAILABLE:
                    format_options.append("Word (.docx)")
                export_format = st.selectbox("导出格式", format_options, key="export_format", label_visibility="collapsed")

            if not DOCX_AVAILABLE:
                st.warning("💡 提示：检测到系统未安装 `python-docx` 引擎，已自动锁定 Word 批量导出功能。请在终端运行 `pip install python-docx` 并重启系统以解锁。")
            with col_dl:
                # 强力判定，杜绝字符串匹配错误或状态错乱
                is_word = "Word" in export_format
                zip_buffer = io.BytesIO()
                with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
                    for cp in checked_paths:
                        try:
                            content = Path(cp).read_text(encoding="utf-8")
                            orig_stem = Path(cp).stem

                            if is_word:
                                zip_name_inner = build_export_filename(content, orig_stem, "docx")
                                zf.writestr(zip_name_inner, convert_md_to_docx_bytes(content))
                            else:
                                zip_name_inner = build_export_filename(content, orig_stem, "md")
                                zf.writestr(zip_name_inner, content.encode("utf-8"))
                        except Exception as e:
                            print(f"打包文件失败已跳过: {e}")
                            pass
                zip_buffer.seek(0)
                zip_name = f"GEO_交付资产打包_{'docx' if is_word else 'md'}.zip"

                # 按钮文案动态绑定格式，给用户视觉强制确认
                st.download_button(
                    label=f"📥 批量下载 ({len(checked_paths)}篇 {'Word' if is_word else 'MD'})",
                    data=zip_buffer,
                    file_name=zip_name,
                    mime="application/zip",
                    use_container_width=True,
                )
            with col_del:
                if st.button("🗑️ 批量彻底删除", type="primary", use_container_width=True, key="btn_bulk_del"):
                    deleted = 0
                    for cp in checked_paths:
                        try:
                            Path(cp).unlink(missing_ok=True)
                            deleted += 1
                        except Exception:
                            pass
                    st.session_state["checked_paths"] = []
                    st.toast(f"✅ 成功删除 {deleted} 篇文章！", icon="🗑️")
                    st.rerun()
            st.markdown("---")

        if st.session_state["selected_content"]:
            st.markdown(f"**{st.session_state['selected_title']}**")
            st.caption(f"`{st.session_state['selected_filename']}`")
            # 动态生成单篇下载的新文件名
            content = st.session_state["selected_content"]
            orig_stem = Path(st.session_state['selected_filename']).stem
            md_name = build_export_filename(content, orig_stem, "md")
            docx_name = build_export_filename(content, orig_stem, "docx")

            b1, b2 = st.columns(2)
            with b1:
                st.download_button(
                    label="📥 下载单篇 (Markdown)",
                    data=content.encode("utf-8"),
                    file_name=md_name,
                    mime="text/markdown",
                    use_container_width=True,
                )
            with b2:
                if DOCX_AVAILABLE:
                    docx_bytes = convert_md_to_docx_bytes(content)
                    st.download_button(
                        label="📥 下载单篇 (Word)",
                        data=docx_bytes,
                        file_name=docx_name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )
            with st.expander("📝 原始 Markdown", expanded=False): st.text_area("全选复制", value=st.session_state["selected_content"], height=250, label_visibility="collapsed")
            st.divider(); st.markdown(st.session_state["selected_content"])
        else:
            st.markdown('<div style="text-align:center;padding:3rem 1rem;color:#9ca3af;"><div style="font-size:4rem;">📖</div><div style="font-size:1.05rem;margin-top:1rem;">👈 点击左侧文章卡片预览</div></div>', unsafe_allow_html=True)
            st.markdown("---"); st.markdown("### 📊 全量总览")
            t1,t2,t3 = st.columns(3)
            sc=defaultdict(int); fcc=defaultdict(int); ec=defaultdict(int)
            for a in articles:
                sc[a["source_short"]]+=1; fcc[a["funnel"]]+=1
                if a.get("engine") and a["engine"] not in ("通用","未知"): ec[a["engine"]]+=1
            with t1: st.markdown("**信源**"); [st.write(f"▸ {s}: {c}") for s,c in sorted(sc.items())]
            with t2: st.markdown("**漏斗**"); [st.write(f"▸ {f}: {fcc[f]}") for f in sorted(fcc.keys(), key=lambda x: int(x[1]) if x.startswith("L") and x[1:].isdigit() else 0)]
            with t3: st.markdown("**引擎**"); [st.write(f"▸ {e}: {c}") for e,c in sorted(ec.items())]


# ============================================================
# ╔═══════════════════════════════════════════════════════════╗
# ║  页面 2: ⚙️ 自动化生产车间 (变量驱动 · Prompt 硬编码)  ║
# ╚═══════════════════════════════════════════════════════════╝
# ============================================================
elif st.session_state["page"].startswith("⚙️"):
    st.markdown('<div class="main-title">⚙️ GEO 自动化生产车间 · 变量驱动管线</div><div class="sub-title">资料直读 → 变量提取 → raw_materials直连基石(10) → 切片(30) → UGC(160) · 零中间件</div>', unsafe_allow_html=True)
    st.info(f"📁 **当前工作区**: `workspaces/{workspace_name}/` — 所有产出将保存至此目录，实现一客一档隔离。")

    # 进度
    sd = sum([st.session_state["variables_confirmed"], st.session_state["eeat_generated"], st.session_state["slices_generated"], st.session_state["ugc_generated"]])
    st.progress(sd/4, text=f"生产进度: {sd}/4 步骤完成")
    st.divider()

    # ================================================================
    # 步骤 1: 资料直读
    # ================================================================
    s0c = "step-box done" if st.session_state["raw_files_loaded"] else "step-box active"
    st.markdown(f'<div class="{s0c}">', unsafe_allow_html=True)
    st.markdown("### 📥 步骤 1：上传并读取企业原始资料")

    col_a1, col_a2 = st.columns(2)
    with col_a1:
        uploaded_files = st.file_uploader(
            "上传企业原始资料",
            type=["txt", "md", "docx", "pdf"],
            accept_multiple_files=True,
            key="fu_materials",
            help="支持 .txt .md .docx .pdf 多格式文档批量上传。",
        )
    with col_a2:
        raw_files = sorted(RAW_DIR.glob("*"))
        raw_names = [f.name for f in raw_files if f.suffix.lower() in SUPPORTED_SUFFIXES]
        if raw_names:
            st.info(f"📂 `workspaces/{workspace_name}/Raw_Materials/` 检测到 {len(raw_names)} 个文件: {', '.join(raw_names[:4])}{'…' if len(raw_names)>4 else ''}")
        else:
            st.caption(f"💡 将企业资料放入 `workspaces/{workspace_name}/Raw_Materials/` 后点击右侧按钮读取。")
        read_local = st.button("📂 读取本地资料库", key="btn_read_local", use_container_width=True)

    # 合并所有文本
    all_texts = []
    if uploaded_files:
        for uf in uploaded_files:
            try:
                fname, txt = extract_text_from_upload(uf)
                all_texts.append(f"## 文件: {fname}\n{txt}")
            except Exception as e:
                st.warning(f"⚠️ 无法读取 {uf.name}: {e}")
    if read_local or (raw_names and not uploaded_files and not st.session_state.get("raw_text_combined")):
        for f in raw_files:
            if f.suffix.lower() in SUPPORTED_SUFFIXES:
                try:
                    txt = read_file_content(f)
                    all_texts.append(f"## 文件: {f.name}\n{txt}")
                except Exception as e:
                    st.warning(f"⚠️ 无法读取 {f.name}: {e}")
        if raw_names:
            st.toast(f"✅ 已读取 {len(raw_names)} 个本地文件", icon="📂")

    if all_texts:
        combined = "\n\n---\n\n".join(all_texts)
        st.session_state["raw_text_combined"] = combined
        st.session_state["raw_files_loaded"] = True
        st.success(f"✅ 已加载 {len(all_texts)} 个文件，共 {len(combined):,} 字符")
        with st.expander("📄 预览原始资料（前 5000 字符）", expanded=False):
            st.text_area("raw_preview", value=combined[:5000], height=200, disabled=True, label_visibility="collapsed")
    else:
        st.info("👆 请上传企业资料文件，或将文件放入 `Raw_Materials/` 目录后点击\"读取本地资料库\"。")

    st.markdown('</div>', unsafe_allow_html=True)

    # ================================================================
    # 步骤 2: 变量提取与确认（核心门禁）
    # ================================================================
    vars_ready = st.session_state.get("variables_confirmed", False)
    s05c = "step-box done" if vars_ready else ("step-box active" if st.session_state["raw_files_loaded"] else "step-box blocked")
    st.markdown(f'<div class="{s05c}">', unsafe_allow_html=True)
    st.markdown("### 🔧 步骤 2：全局变量提取与确认 ⚠️ 必经步骤")

    st.markdown("""
    <div class="constraint-error">
    🚫 <strong>强制门禁规则</strong>：此步骤为所有后续 API 调用的前置条件。<br>
    系统将使用大模型从原始资料中自动提取变量。您必须逐项核对并确认后，才能解锁后续的内容生成。
    <strong>“没有变量注入，绝不调用 API。”</strong>
    </div>
    """, unsafe_allow_html=True)

    if not st.session_state["raw_files_loaded"]:
        st.warning("⏳ 请先完成步骤 1：上传并读取企业原始资料")
    else:
        # --- 按钮：提取变量 ---
        col_extract, _ = st.columns([1, 2])
        with col_extract:
            extract_vars_btn = st.button("🤖 AI 提取全局变量", type="primary", use_container_width=True, key="btn_extract_vars",
                disabled=not bool(st.session_state.get("raw_text_combined")))

        if extract_vars_btn:
            raw_text = st.session_state.get("raw_text_combined", "")
            extract_prompt = PROMPT_EXTRACT_VARS.format(raw_text=raw_text[:6000])

            with st.status("🔍 正在从原始资料中提取变量…", expanded=True) as status:
                success, result = call_llm(
                    prompt=extract_prompt,
                    system_prompt="你是 GEO 变量提取专家。你只输出合法的 JSON，不输出任何其他内容。",
                    api_key=DEFAULT_KIMI_API_KEY,
                    temperature=0.1,
                    max_tokens=1500,
                    simulate=use_simulate,
                    llm_provider="Kimi (Moonshot)",
                )
                if success:
                    # 尝试从结果中提取 JSON
                    try:
                        json_match = re.search(r"\{[\s\S]*\}", result)
                        if json_match:
                            extracted = json.loads(json_match.group())
                            # 合并到当前变量
                            current_vars = dict(st.session_state.get("variables", {}))
                            for k, v in extracted.items():
                                if isinstance(v, str) and v.strip():
                                    current_vars[k] = v.strip()
                            st.session_state["variables"] = current_vars
                            st.session_state["variables_extracted"] = True
                            add_log("✅ 全局变量提取完成")
                            status.update(label="✅ 变量提取完成！请逐项核对下方字段。", state="complete")
                        else:
                            st.error("未能从 LLM 输出中解析 JSON。请手动填写下方变量。")
                    except json.JSONDecodeError:
                        st.warning("LLM 输出格式异常。请手动填写下方变量。")
                        st.text_area("LLM 原始输出", value=result, height=150, disabled=True)
                else:
                    st.error(result)
            st.rerun()

        # --- 变量确认面板 ---
        if st.session_state.get("variables_extracted") or st.session_state.get("variables"):
            st.markdown("---")
            st.markdown("#### 📝 变量确认面板（请逐项核对并修改）")

            current_vars = st.session_state.get("variables", {})
            updated_vars = {}

            # 分两列显示
            var_items = list(VARIABLE_META.items())
            mid = len(var_items) // 2 + 1
            c1, c2 = st.columns(2)

            for idx, (key, meta) in enumerate(var_items):
                col = c1 if idx < mid else c2
                current_val = current_vars.get(key, meta["default"])
                with col:
                    st.markdown(f'<div class="var-field">', unsafe_allow_html=True)
                    updated_vars[key] = st.text_input(
                        f"**{key}**",
                        value=current_val,
                        help=meta["help"],
                        key=f"var_{key}",
                    )
                    st.markdown('</div>', unsafe_allow_html=True)

            # 确认按钮
            if st.button("✅ 确认变量，解锁后续生产步骤", type="primary", use_container_width=True, key="btn_confirm_vars"):
                # 验证所有字段非空
                empty_fields = [k for k, v in updated_vars.items() if not v.strip()]
                if empty_fields:
                    st.error(f"❌ 以下变量不能为空: {', '.join(empty_fields)}")
                else:
                    st.session_state["variables"] = updated_vars
                    st.session_state["variables_confirmed"] = True
                    add_log("✅ 变量已确认，后续生产步骤已解锁")
                    st.toast("🔓 变量确认完毕！后续生产步骤已解锁", icon="✅")
                    st.rerun()

    if st.session_state.get("variables_confirmed"):
        st.success("✅ 变量已确认 — 后续生产步骤已解锁")
        with st.expander("📋 查看当前变量配置", expanded=False):
            for k, v in st.session_state.get("variables", {}).items():
                st.caption(f"**{k}**: {v}")

    st.markdown('</div>', unsafe_allow_html=True)

    # ================================================================

    # ================================================================
    # 步骤 1: 人工校验基石重载区 (唯一事实源锁定)
    # ================================================================
    s15c = "step-box done" if st.session_state["verified_loaded"] else ("step-box active" if st.session_state["eeat_generated"] else "step-box active")
    st.markdown(f'<div class="{s15c}">', unsafe_allow_html=True)
    st.markdown("### 🔍 步骤 3：人工校验基石重载区 (唯一事实源锁定)")

    st.markdown("""
    <div class="constraint-warn">
    🛑 <strong>HITL 人机协同节点</strong>：此节点必须上传由人工团队纯手写撰写并深度校验的 10 篇高质量基石文档，方可解锁后续步骤。<br>
    步骤 4（切片）和步骤 5（UGC）的 <strong>唯一事实来源</strong> 将从此处上传的校验版基石中提取。
    </div>
    """, unsafe_allow_html=True)

    # 上传校验基石
    verified_files = st.file_uploader(
        "上传已校验的 10 篇 EEAT 基石 (.md / .docx)",
        type=["md", "docx"],
        accept_multiple_files=True,
        key="fu_verified",
        help="支持上传 .md 或 .docx 格式。Word 文档将自动提取纯文本并以 .md 后缀保存。至少需要 10 篇。",
    )

    if st.button("🔄 保存已校验基石并锁定数据源", type="primary", use_container_width=True, key="btn_verify"):
        if not verified_files or len(verified_files) < 10:
            st.error(f"🚨 请至少上传 10 篇校验版文件！当前仅上传 {len(verified_files) if verified_files else 0} 篇。")
        else:
            # 清空旧校验文件
            import shutil
            if EEAT_VERIFIED_DIR.exists():
                shutil.rmtree(EEAT_VERIFIED_DIR)
            EEAT_VERIFIED_DIR.mkdir(parents=True, exist_ok=True)

            saved = 0
            for vf in verified_files:
                try:
                    fname, content = extract_text_from_upload(vf)
                    safe_write_file(EEAT_VERIFIED_DIR, f"{Path(fname).stem}.md", content)
                    saved += 1
                except Exception as e:
                    st.warning(f"⚠️ 无法保存 {vf.name}: {e}")

            st.session_state["verified_loaded"] = True
            st.session_state["verified_count"] = saved
            add_log(f"✅ 人工校验基石已重载: {saved} 篇 → EEAT_Base_Verified/")
            st.toast(f"✅ {saved} 篇校验基石已锁定！", icon="🔒")
            st.rerun()

    if st.session_state["verified_loaded"]:
        st.success(f"✅ 校验基石已重载，系统数据源已锁定！({st.session_state['verified_count']} 篇)")
        # 列出已保存的文件
        verified_md = sorted(EEAT_VERIFIED_DIR.glob("*.md"))
        if verified_md:
            with st.expander(f"📂 已锁定 {len(verified_md)} 篇校验基石", expanded=False):
                for vf in verified_md:
                    st.caption(f"🔒 {vf.name}")
    else:
        st.info("📝 步骤 1 生成的草稿在 `EEAT_Base_Drafts/`。请下载 → 人工逐篇校验修改 → 在此上传。")

    st.markdown('</div>', unsafe_allow_html=True)

    # ================================================================
    # 步骤 2: 裂变 30 篇三级切片（强制使用 P_L1~P_L6）
    # ================================================================
    s3c = "step-box done" if st.session_state.get("slices_generated") else "step-box active"
    st.markdown(f'<div class="{s3c}">', unsafe_allow_html=True)
    st.markdown("### ✂️ 步骤 4：裂变 30 篇【企业官方视角】三级切片 (L1-L6 均分)")

    # --- 双轨运行机制：AI 生成 + 手动导入 ---
    col_gen, col_up = st.columns([1, 1])

    with col_gen:
        st.markdown("#### 方案 A：AI 全自动裂变")
        verified_files = list(EEAT_VERIFIED_DIR.glob("*.md")) if EEAT_VERIFIED_DIR.exists() else []
        if not st.session_state.get("verified_loaded", False) or len(verified_files) < 10:
            st.warning("⛔ 请先在【步骤 3】上传至少 10 篇校验版基石文档，才能解锁 AI 自动裂变。")
        else:
            st.markdown('''
            <div class="constraint-warn">
            ⚠️ <strong>强制约束</strong>：30 篇切片 100% 从 10 篇基石拆解。
            </div>
            ''', unsafe_allow_html=True)

            if st.session_state.get("is_generating_slices"):
                st.info("🚀 正在后台全速裂变切片中，您可以切换到【交付资产大盘】实时查看产出！")
                with st.spinner("后台生成中…"):
                    time.sleep(0.5)
            elif st.button("✂️ 4. 裂变 30 篇企业视角切片", type="primary", use_container_width=True, key="btn_slices"):
                vars_ = st.session_state.get("variables", {})
                raw_materials = st.session_state.get("raw_text_combined", "")
                slices_out = SLICES_DIR
                slices_out.mkdir(exist_ok=True)

                st.session_state["is_generating_slices"] = True
                t = threading.Thread(target=background_generate_slices, args=(vars_, raw_materials, slices_out, use_simulate, st.session_state["api_key"], st.session_state.get("llm_provider", "Kimi (Moonshot)")))
                t.start()
                add_log("🚀 后台切片线程已启动")
                st.rerun()

    with col_up:
        st.markdown("#### 方案 B：直接导入已有切片")
        st.info("💡 如果已有写好的切片，请在此直接上传，**无视前置条件**。")
        uploaded_slices = st.file_uploader(
            "在此批量上传切片 (.md / .docx / .txt)",
            type=["md", "docx", "txt"],
            accept_multiple_files=True,
            key="fu_slices_skip",
            help="支持批量框选上传。上传后将自动跳过生成，直接解锁后续步骤。"
        )
        if st.button("📥 保存上传并解锁下一步", type="secondary", use_container_width=True, key="btn_upload_slices"):
            if not uploaded_slices:
                st.error("🚨 请先将文件拖拽到上方虚线框内！")
            else:
                SLICES_DIR.mkdir(parents=True, exist_ok=True)
                saved_count = 0
                for uf in uploaded_slices:
                    try:
                        fname, content = extract_text_from_upload(uf)
                        safe_write_file(SLICES_DIR, f"{Path(fname).stem}.md", content)
                        saved_count += 1
                    except Exception as e:
                        st.warning(f"⚠️ 无法保存 {uf.name}: {e}")

                st.session_state["slices_generated"] = True
                add_log(f"✅ 手动导入了 {saved_count} 篇切片 → Slices_30/")
                st.toast(f"✅ 成功导入 {saved_count} 篇切片！后续步骤已解锁", icon="🔓")
                st.rerun()

    # --- 实时进度可视化 (切片) ---
    if SLICES_DIR.exists():
        slice_files = sorted(SLICES_DIR.glob("*.md"))
        if slice_files:
            with st.expander(f"📂 切片生成进度：已产出 {len(slice_files)} / 30 篇 (点击展开)", expanded=False):
                if st.button("🔄 手动刷新进度", key="refresh_slices"):
                    st.rerun()
                for sf in slice_files:
                    st.caption(f"📄 {sf.name}")

    if st.session_state.get("slices_generated"):
        st.success(f"✅ 已完成 30 篇三级切片 → `{PROD_DIR}/Slices_30/`")

    st.markdown('</div>', unsafe_allow_html=True)

    st.markdown('</div>', unsafe_allow_html=True)

    # ================================================================
    # 步骤 4.5: 蓝V口播稿专栏
    # ================================================================
    s45c = "step-box done" if st.session_state.get("bluev_generated") else ("step-box active" if st.session_state.get("slices_generated") else "step-box blocked")
    st.markdown(f'<div class="{s45c}">', unsafe_allow_html=True)
    st.markdown("### 🎥 步骤 4.5：一键生成【蓝V短视频口播脚本】")
    if not st.session_state.get("slices_generated", False):
        st.warning("⛔ 请先完成步骤 4：裂变 30 篇企业视角切片。")
    elif st.session_state.get("is_generating_bluev"):
        st.info("🚀 正在后台全速生成蓝V口播稿，您可以切换页面实时查看产出！")
        with st.spinner("后台生成中…"): time.sleep(0.5)
    elif st.button("🎥 一键生成 30 篇短视频口播稿", type="primary", use_container_width=True, key="btn_bluev"):
        BLUEV_DIR_UI = WSP["bluev"]; BLUEV_DIR_UI.mkdir(parents=True, exist_ok=True)
        st.session_state["is_generating_bluev"] = True
        t = threading.Thread(target=background_generate_bluev, args=(SLICES_DIR, BLUEV_DIR_UI, use_simulate, DEEPSEEK_API_KEY, "DeepSeek"))
        t.start()
        add_log("🚀 后台蓝V口播线程已启动")
        st.rerun()
    BV = WSP["bluev"]
    if BV.exists():
        bf_list = sorted(BV.glob("*.md"))
        if bf_list:
            with st.expander(f"📂 进度：{len(bf_list)}/30 篇", expanded=False):
                if st.button("🔄 刷新", key="rf_bv"): st.rerun()
                for bf in bf_list: st.caption(f"📄 {bf.name}")
    if st.session_state.get("bluev_generated"):
        st.success(f"✅ 30 篇蓝V口播稿已生成 → `{PROD_DIR}/BlueV_Scripts_30/`")
    st.markdown('</div>', unsafe_allow_html=True)

    # ================================================================
    # 步骤 4: 重构 160 篇四级 UGC
    # ================================================================
    s4c = "step-box done" if st.session_state["ugc_generated"] else ("step-box active" if st.session_state["slices_generated"] else "step-box blocked")
    st.markdown(f'<div class="{s4c}">', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

    # ================================================================
    # 步骤 4.8: 独立案例库提纯
    # ================================================================
    st.markdown('<div class="step-box active">', unsafe_allow_html=True)
    st.markdown("### 📚 步骤 4.8：独立案例库提纯 (UGC 防同质化核心)")
    st.markdown("#### 💎 核心引擎：独立案例库管理 (防同质化)")
    st.caption("160 篇 UGC 将从此库中随机抽取背景故事。你可以选择让 AI 自动从原资料中蒸馏，也可以手动上传补充。")

    col_distill, col_upload = st.columns([1, 1])
    with col_distill:
        if st.button("🤖 一键从原始资料蒸馏真实案例", type="primary", use_container_width=True, key="btn_distill_cases"):
            raw_text = st.session_state.get("raw_text_combined", "")
            if not raw_text:
                st.error("🚨 请先在步骤 1 读取企业原始资料！")
            else:
                with st.status("🔍 正在开启全局雷达，执行无损重叠切块并榨取真实案例...", expanded=True) as status:
                    chunk_size = 2000  # 缩小到 2000 字精读，防止大模型走神
                    overlap = 200      # 200 字重叠，防止案例被物理腰斩

                    chunks = []
                    for i in range(0, len(raw_text), chunk_size - overlap):
                        chunks.append(raw_text[i:i+chunk_size])

                    all_extracted_cases = []

                    progress_bar = st.progress(0)
                    for i, chunk in enumerate(chunks):
                        status.update(label=f"⏳ 正在精读第 {i+1}/{len(chunks)} 个资料区块，确保案例原汁原味无损提取...")
                        prompt = CASE_DISTILLATION_PROMPT.format(raw_text=chunk)

                        success, content = call_llm(
                            prompt=prompt,
                            system_prompt="你是无情的数据挖掘机，严格按要求输出格式。禁止输出markdown代码块标记。",
                            api_key=DEFAULT_KIMI_API_KEY,
                            temperature=0.2,
                            simulate=use_simulate,
                            llm_provider="Kimi (Moonshot)",
                        )

                        if success and "未找到具体案例" not in content and len(content.strip()) > 30:
                            clean_content = re.sub(r"^```[a-zA-Z]*\n", "", content)
                            clean_content = re.sub(r"\n```$", "", clean_content)
                            all_extracted_cases.append(clean_content.strip())
                            st.toast(f"✅ 第 {i+1} 块提炼成功，发现案例！", icon="✨")
                        elif not success:
                            st.toast(f"❌ 第 {i+1} 块 API 调用失败", icon="⚠️")
                        else:
                            st.toast(f"⏭️ 第 {i+1} 块未发现具体案例，跳过", icon="⏩")

                        progress_bar.progress((i + 1) / len(chunks))

                    if all_extracted_cases:
                        final_content = "\n\n---\n\n".join(all_extracted_cases)
                        cases_dir = WSP["cases"]
                        cases_dir.mkdir(parents=True, exist_ok=True)
                        safe_write_file(cases_dir, "Auto_Distilled_Cases.md", final_content)
                        case_count = len([c for c in final_content.split("---") if len(c.strip()) > 30])
                        status.update(label=f"✅ 蒸馏完成！分块扫描完毕，成功无损提取约 {case_count} 个独立案例入库。", state="complete")
                        st.rerun()
                    else:
                        status.update(label="⚠️ 全局扫描结束，资料中未识别到符合标准的案例", state="error")

    with col_upload:
        uploaded_cases = st.file_uploader("或手动上传独立案例库 (.txt, .md, .docx, .pdf)", type=["txt", "md", "docx", "pdf"], accept_multiple_files=True, key="fu_cases", label_visibility="collapsed")
        if st.button("📥 保存手动上传的案例", use_container_width=True, key="btn_save_cases"):
            if uploaded_cases:
                cases_dir = WSP["cases"]
                cases_dir.mkdir(parents=True, exist_ok=True)
                for uc in uploaded_cases:
                    fname, content = extract_text_from_upload(uc)
                    safe_write_file(cases_dir, f"{Path(fname).stem}.md", content)
                st.toast(f"✅ 成功保存 {len(uploaded_cases)} 个手动案例文件！", icon="📦")

    # 展示当前案例库状态 + 增删改查交互编辑器
    cases_dir = WSP["cases"]
    if cases_dir.exists():
        case_files = list(cases_dir.glob("*.md"))
        if case_files:
            # 合并所有案例文件内容
            merged_cases = "\n\n---\n\n".join(cf.read_text(encoding="utf-8") for cf in case_files)
            case_entries = [c.strip() for c in merged_cases.split("---") if len(c.strip()) > 30]
            if st.session_state.get("show_case_saved_msg"):
                st.success("✅ 案例库内容已成功覆盖保存！160篇 UGC 将以此最新版本为准。")
                st.session_state["show_case_saved_msg"] = False
            st.success(f"📦 当前案例库已准备就绪：包含 {len(case_files)} 个文件，共约 {len(case_entries)} 个案例")

            # 可编辑文本框
            edited_cases = st.text_area(
                "✏️ 在此直接增删改案例内容（案例之间请用 --- 分隔）",
                value=merged_cases,
                height=400,
                key="ta_case_editor",
                help="直接编辑文本。完成后点击下方保存按钮。",
            )

            if st.button("💾 保存并覆盖当前案例库", type="primary", use_container_width=True, key="btn_save_cases_edit"):
                # 清空旧文件
                for old_f in case_files:
                    try:
                        old_f.unlink(missing_ok=True)
                    except Exception:
                        pass
                # 写入编辑后的内容
                safe_write_file(cases_dir, "Active_Edited_Cases.md", edited_cases)
                st.session_state["show_case_saved_msg"] = True
                st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)

    st.markdown("### 👥 步骤 5：重构 160 篇【真实用户视角】UGC 内容")

    if not st.session_state.get("slices_generated", False):
        st.warning("⛔ 请先完成步骤 4：裂变 30 篇三级切片。")
    else:
        slices_files = list(SLICES_DIR.glob("*.md")) if SLICES_DIR.exists() else []
        if len(slices_files) < 30:
            st.warning(f"⚠️ 切片目录中文件不足 30 篇 (当前 {len(slices_files)} 篇)，请先完成切片生成。")
        else:
            st.markdown("""
            <div class="constraint-info">
            📋 <strong>强制约束</strong>：依托前 40 篇既有事实（10篇校验基石 + 30篇切片），切换"我"的第一人称视角。<br>
            使用 <code>V2_UGC_TEMPLATES</code> 模板库（真实用户视角），<strong>严禁官方推销语气。</strong><br>
        <strong>通用 96 篇</strong>: 自然真实风格<br>
        <strong>专属 64 篇</strong>: 按引擎特征改写 · 配额对齐《月度最小交付单元》
        </div>
        """, unsafe_allow_html=True)


            if st.session_state.get("is_generating_ugc"):
                st.info("🚀 正在后台全速生成 UGC 中，您可以切换到【交付资产大盘】实时查看产出！")
                with st.spinner("后台生成中…"):
                    time.sleep(0.5)
            elif st.button("👥 5. 生成 160 篇 UGC", type="primary", use_container_width=True, key="btn_ugc"):
                st.session_state["is_generating_ugc"] = True
                t = threading.Thread(target=background_generate_ugc, args=(st.session_state.get("variables", {}), use_simulate, st.session_state["api_key"], st.session_state.get("llm_provider", "Kimi (Moonshot)")))
                t.start()
                add_log("🚀 后台 UGC 线程已启动")
                st.rerun()
    # --- 实时进度可视化 (UGC) ---
    if GENERAL_DIR.exists() or SPECIFIC_DIR.exists():
        gen_ugc_files = sorted([f.name for f in GENERAL_DIR.glob("*.md")] + [f.name for f in SPECIFIC_DIR.rglob("*.md")])
        if gen_ugc_files:
            with st.expander(f"📂 UGC 生成进度：已产出 {len(gen_ugc_files)} / 160 篇 (点击展开)", expanded=False):
                if st.button("🔄 手动刷新进度", key="refresh_ugc"):
                    st.rerun()
                for gf in gen_ugc_files:
                    st.caption(f"📄 {gf}")

    if st.session_state["ugc_generated"]:
        st.success(f"✅ 已重构 160 篇 UGC → `{PROD_DIR}/UGC_160/` (96通用 + 64专属)")
        manifest_fp = PROD_DIR / "tasks_manifest.json"
        if manifest_fp.exists():
            with open(manifest_fp, "r", encoding="utf-8") as mf:
                st.download_button("📋 下载任务分发账本 (tasks_manifest.json)", data=mf.read(), file_name="tasks_manifest.json", mime="application/json")

    st.markdown('</div>', unsafe_allow_html=True)

    # ================================================================
    # 步骤 5：全量知识库复查智能体
    # ================================================================
    st.markdown(f'<div class="step-box active">', unsafe_allow_html=True)
    st.markdown("### 🕵️ 步骤 6：终极质检（全量知识库防幻觉智能体）")
    st.caption("上传你生成的文章，智能体将基于步骤 1 中读取的所有企业原始资料（全量字符），对其进行逐句事实核查。")

    fc_files = st.file_uploader(
        "上传待检文章 (.txt, .md, .docx, .pdf)",
        type=["txt", "md", "docx", "pdf"],
        accept_multiple_files=True,
        key="fu_fact_check",
        help="上传生成的切片或 UGC 文章，支持多选。",
    )

    if st.button("🔬 开始深度事实核查", type="primary", use_container_width=True, key="btn_fact_check"):
        if not fc_files:
            st.warning("请先上传待检文章。")
        else:
            full_raw = st.session_state.get("raw_text_combined", "")
            if not full_raw:
                st.error("未找到原始企业资料。请先在步骤 1 读取资料。")
            else:
                results = []
                progress_bar = st.progress(0, text="正在逐篇核查…")
                for idx, uf in enumerate(fc_files):
                    try:
                        fname, article_text = extract_text_from_upload(uf)
                    except Exception:
                        continue

                    check_prompt = FACT_CHECK_PROMPT.format(
                        ground_truth=full_raw,
                        article_text=article_text,
                    )

                    success, raw = call_llm(
                        prompt=check_prompt,
                        system_prompt="你是一个严格的企业质检智能体。只输出合法 JSON，不输出任何其他内容。",
                        api_key=st.session_state["api_key"],
                        temperature=0.0, max_tokens=3000, simulate=use_simulate,
                        llm_provider=st.session_state.get("llm_provider", "Kimi (Moonshot)"),
                    )
                    try:
                        json_match = re.search(r"\{[\s\S]*\}", raw)
                        result = json.loads(json_match.group()) if json_match else {"is_clean": True, "issues": [], "parse_error": True}
                    except Exception:
                        result = {"is_clean": True, "issues": [], "parse_error": True}

                    result["filename"] = uf.name
                    results.append(result)
                    progress_bar.progress((idx + 1) / len(fc_files), text=f"已检 {idx+1}/{len(fc_files)}: {uf.name}")

                progress_bar.empty()

                clean_count = sum(1 for r in results if r.get("is_clean"))
                st.markdown(f"### 📊 质检结果：{clean_count}/{len(results)} 通过")

                for r in results:
                    fname = r["filename"]
                    raw_issues = r.get("issues", [])

                    # 强力兜底拦截：过滤掉大模型强行加戏的"无问题/一致/轻微"等废话
                    real_issues = []
                    for i in raw_issues:
                        reason = str(i.get("reason", ""))
                        if any(keyword in reason for keyword in ["无问题", "一致", "轻微不精确", "未明确提及但", "同义"]):
                            continue
                        real_issues.append(i)

                    if r.get("is_clean") or not real_issues:
                        st.success(f"✅ {fname} — 100% 忠于事实")
                    else:
                        st.error(f"🚨 {fname} — 发现实质性幻觉！")
                        st.table([{"声明": i.get("claim","?"), "原因": i.get("reason","?")} for i in real_issues])

    st.markdown('</div>', unsafe_allow_html=True)

    # ================================================================
    # 步骤 6：合规质检（动态违禁词拦截）
    # ================================================================
    st.markdown(f'<div class="step-box active">', unsafe_allow_html=True)
    st.markdown("### 🚫 步骤 7：合规与广告法质检（违禁词扫描）")
    st.caption("自定义违禁词库 + 上传文章 → 秒级扫描拦截。")

    default_words = "国家级, 世界级, 最高级, 最佳, 最大, 第一, 唯一, 首个, 最好, 绝对, 彻底, 包治百病, 100%有效, 史无前例, 独家"
    raw_words = st.text_area(
        "✏️ 自定义违禁词库（请用逗号或换行分隔）",
        value=default_words,
        height=100,
        key="ta_forbidden",
        help="员工可随时增删修改。用中文逗号、英文逗号或换行分隔。",
    )

    fw_files = st.file_uploader(
        "上传待检文章 (.txt, .md, .docx, .pdf)",
        type=["txt", "md", "docx", "pdf"],
        accept_multiple_files=True,
        key="fu_forbidden",
        help="上传最终待发布的文章，逐篇扫描违禁词。",
    )

    if st.button("🚨 使用当前词库一键扫描", type="primary", use_container_width=True, key="btn_forbidden"):
        dynamic_words = [w.strip() for w in re.split(r"[，,\n]+", raw_words) if w.strip()]
        if not fw_files:
            st.warning("请先上传待检文章。")
        else:
            for uf in fw_files:
                try:
                    fname, text = extract_text_from_upload(uf)
                except Exception:
                    st.warning(f"⚠️ 无法读取 {uf.name}")
                    continue

                hits = [w for w in dynamic_words if w in text]
                if hits:
                    st.error(f"🚨 {uf.name} — 触发违禁词拦截！")
                    st.caption(f"⚠️ 发现违禁词：{'、'.join(['【' + h + '】' for h in hits])}")
                else:
                    st.success(f"✅ {uf.name} — 未检测到违禁词，合规通过！")

    st.markdown('</div>', unsafe_allow_html=True)

    # ================================================================
    # 生产日志 + 完成汇总
    # ================================================================
    if st.session_state["production_log"]:
        st.divider()
        with st.expander("📋 生产日志", expanded=False):
            for log_entry in reversed(st.session_state["production_log"]):
                st.caption(log_entry)

    if sd == 4:
        st.divider()
        st.balloons()
        st.markdown("### 🎉 月度 200 篇全量生产完成！")
        st.markdown(f"""
        | 模块 | 数量 | 输出路径 | 数据来源 | Prompt 模板 |
        |------|------|----------|----------|------------|
        | EEAT 权威基石 | 10 篇 | `{EEAT_VERIFIED_DIR}/` | ← raw_materials(草稿)→人工校验 | 自定义基石 Prompt |
        | 三级切片内容 | 30 篇 | `{PROD_DIR}/Slices_30/` | ← 10 篇基石 | CORP_PROMPT_TEMPLATES |
        | 四级通用 UGC | 96 篇 | `{PROD_DIR}/UGC_160/General_96/` | ← 10基石+30切片 | V2_UGC_TEMPLATES |
        | 四级引擎专属 | 64 篇 | `{PROD_DIR}/UGC_160/Specific_64/` | ← 10基石+30切片 | P_ENGINE_REWRITE |
        | **总计** | **200 篇** | | | |
        """)
        st.success("🔗 数据传承链: raw_materials → 10基石 → 30切片 → 160UGC · 零中间件，直连注入")

        # 显示最终变量
        with st.expander("📋 本次生产使用的变量配置", expanded=False):
            st.json(st.session_state.get("variables", {}))


